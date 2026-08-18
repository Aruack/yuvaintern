"""
Generate Publication-Grade Microsoft Word (.docx) Report
Week 1 Task: Data Acquisition, Cleaning, and Preprocessing
Author: Aryan Kumar (Virtual Data Science Trainee)
"""

import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

DOC_PATH = "Week_1_Data_Acquisition_Cleaning_Report.docx"
ASSETS_DIR = "assets"

def set_cell_background(cell, fill_color):
    """Sets background color of a table cell (e.g. '1A365D')."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding/margins for table cells (in twips, 20 twips = 1 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout_box(doc, title, text, box_type="note"):
    """
    Creates a styled callout box with a colored left accent border.
    box_type: 'note' (blue), 'challenge' (amber), 'success' (green), 'warning' (red)
    """
    colors = {
        "note": {"border": "2B6CB0", "bg": "EBF8FF", "title": "2B6CB0"},
        "challenge": {"border": "DD6B20", "bg": "FFFAF0", "title": "C05621"},
        "success": {"border": "38A169", "bg": "F0FFF4", "title": "276749"},
        "warning": {"border": "E53E3E", "bg": "FFF5F5", "title": "9B2C2C"}
    }
    cfg = colors.get(box_type, colors["note"])
    
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, cfg["bg"])
    set_cell_margins(cell, top=140, bottom=140, left=200, right=180)
    
    # Set left border thick, clear other borders
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="36" w:space="0" w:color="{cfg['border']}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run_title = p.add_run(f"📌 {title}\n")
    run_title.bold = True
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(10.5)
    run_title.font.color.rgb = RGBColor.from_string(cfg["title"])
    
    run_text = p.add_run(text)
    run_text.font.name = "Calibri"
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(45, 55, 72)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_code_block(doc, code_str, caption=None):
    """Adds a syntax-highlighted code block container."""
    if caption:
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(6)
        p_cap.paragraph_format.space_after = Pt(2)
        r_cap = p_cap.add_run(f"Code Snippet: {caption}")
        r_cap.font.name = "Calibri"
        r_cap.font.size = Pt(9.5)
        r_cap.font.bold = True
        r_cap.font.color.rgb = RGBColor(74, 85, 104)
        
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F7FAFC")
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    
    # Border: subtle gray around
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="6" w:space="0" w:color="CBD5E0"/>
            <w:left w:val="single" w:sz="18" w:space="0" w:color="4A5568"/>
            <w:bottom w:val="single" w:sz="6" w:space="0" w:color="CBD5E0"/>
            <w:right w:val="single" w:sz="6" w:space="0" w:color="CBD5E0"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(code_str)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(26, 32, 44)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_heading_styled(doc, text, level):
    """Adds properly colored and spaced heading."""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.keep_with_next = True
    run = h.runs[0]
    run.font.name = "Calibri"
    
    if level == 1:
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(26, 54, 93) # Deep Navy (#1A365D)
    elif level == 2:
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(43, 108, 176) # Slate Blue (#2B6CB0)
    elif level == 3:
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(2)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(45, 55, 72) # Charcoal
    return h

def add_figure_with_caption(doc, img_filename, caption_text, width_inches=6.2):
    """Embeds an image from assets and attaches a styled caption below it."""
    img_path = os.path.join(ASSETS_DIR, img_filename)
    if not os.path.exists(img_path):
        print(f"Warning: Image {img_path} not found.")
        return
        
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(2)
    p_img.paragraph_format.keep_with_next = True
    
    run_img = p_img.add_run()
    run_img.add_picture(img_path, width=Inches(width_inches))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(0)
    p_cap.paragraph_format.space_after = Pt(8)
    
    r_cap = p_cap.add_run(f"Figure: {caption_text}")
    r_cap.font.name = "Calibri"
    r_cap.font.size = Pt(9)
    r_cap.font.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

def style_table(table, col_widths, headers, data, align_right_cols=[]):
    """Styles a table with professional corporate headers, alternating rows, and border padding."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Set headers
    hdr_row = table.rows[0]
    hdr_tr = hdr_row._tr.get_or_add_trPr()
    hdr_tr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    
    for i, title in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = title
        set_cell_background(cell, "1A365D")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i in align_right_cols else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for r in p.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            
    # Add Data Rows
    for row_idx, row_data in enumerate(data):
        row = table.add_row()
        bg_color = "F7FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(val)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=70, bottom=70, left=100, right=100)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if col_idx in align_right_cols else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.font.name = "Calibri"
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(45, 55, 72)
                
    # Apply widths
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Inches(w)
            
    doc_add_spacing = docx.text.paragraph.Paragraph(table._element, table._parent)

def build_report():
    doc = Document()
    
    # Page setup - 0.8 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Document Title Block / Cover Header
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(2)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_main = title_p.add_run("WEEK 1 TECHNICAL REPORT\nDATA ACQUISITION, CLEANING, AND PREPROCESSING")
    r_main.font.name = "Calibri"
    r_main.font.size = Pt(20)
    r_main.font.bold = True
    r_main.font.color.rgb = RGBColor(26, 54, 93)
    
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(2)
    sub_p.paragraph_format.space_after = Pt(12)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = sub_p.add_run("Virtual Data Science with Python Trainee Program | YuvaInternship")
    r_sub.font.name = "Calibri"
    r_sub.font.size = Pt(12)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(74, 85, 104)
    
    # Metadata Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_widths = [2.2, 4.3]
    meta_info = [
        ("Trainee Name / Author", "Aryan Kumar"),
        ("Program / Role", "Virtual Data Science with Python Trainee"),
        ("Core Stack", "Python 3.13 | Pandas | NumPy | Scikit-Learn | Matplotlib | Seaborn"),
        ("Submission Date & Status", "Week 1 Milestone | Comprehensive Final Deliverable")
    ]
    for idx, (lbl, val) in enumerate(meta_info):
        row = meta_table.rows[idx]
        c1, c2 = row.cells[0], row.cells[1]
        c1.text = lbl
        c2.text = val
        set_cell_background(c1, "EDF2F7")
        set_cell_background(c2, "FFFFFF")
        set_cell_margins(c1, 50, 50, 80, 80)
        set_cell_margins(c2, 50, 50, 80, 80)
        c1.paragraphs[0].runs[0].font.bold = True
        c1.paragraphs[0].runs[0].font.size = Pt(9.5)
        c2.paragraphs[0].runs[0].font.size = Pt(9.5)
        c1.width = Inches(meta_widths[0])
        c2.width = Inches(meta_widths[1])
        
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # Section 1: Executive Summary
    add_heading_styled(doc, "1. Executive Summary & Project Objectives", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Data quality is the cornerstone of any production-grade analytical and predictive machine learning system. "
        "Raw datasets harvested from real-world systems, transactional databases, and legacy CRMs are inherently contaminated with "
        "anomalies, missing observations, malformed strings, erroneous sensor/financial values, and severe distributional skewness. "
        "The objective of this Week 1 milestone is to establish a rigorous, repeatable, and scalable end-to-end Python pipeline "
        "that ingests a publicly available benchmark dataset (Customer Intelligence and Churn Analytics), performs holistic data "
        "auditing, executes multi-stage data remediation, and transforms the raw records into an optimized, model-ready format."
    )
    
    add_callout_box(
        doc,
        "Core Milestone Objectives",
        "• Data Acquisition & Ingestion: Programmatically load and audit raw transactional data.\n"
        "• Data Quality Auditing: Detect missingness mechanisms (MCAR/MAR), schema anomalies, and invalid data types.\n"
        "• Advanced Remediation: Execute KNN numerical imputation, categorical mode imputation, string normalization, and domain boundary constraints.\n"
        "• Outlier Diagnostics: Perform IQR Winsorization and assess multi-dimensional anomaly behavior.\n"
        "• Skewness & Feature Engineering: Apply Log1p transformations, binning, service density indicators, and robust scaling.",
        "note"
    )
    
    # Section 2: Architecture Flowchart
    add_heading_styled(doc, "2. End-to-End Data Preprocessing Architecture", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "The preprocessing architecture follows an industrial 5-stage ETL & Feature Engineering paradigm. "
        "Each stage encapsulates strict validation gates ensuring zero data leakage and maintaining statistical fidelity."
    )
    add_figure_with_caption(doc, "08_data_cleaning_pipeline_flowchart.png", "Industrial 5-Stage Data Cleaning & Preprocessing Architectural Flowchart")
    
    # Section 3: Data Acquisition & Structural Audit
    add_heading_styled(doc, "3. Data Acquisition & Initial Exploration", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "The dataset represents a customer behavioral database encompassing demographic profiles, subscribed services, contract details, "
        "billing metrics, and churn status across 2,500 observation records and 20 distinct feature attributes. "
        "During initial ingestion, programmatic inspection identified widespread structural data quality defects:"
    )
    
    # Table of Features
    schema_headers = ["Attribute", "Raw Type", "Target Domain", "Observed Defect / Noise Pattern"]
    schema_data = [
        ["CustomerID", "object", "Alphanumeric ID", "Duplicate customer primary keys (~20 duplicate entries)"],
        ["Gender", "object", "Binary (Male/Female)", "Inconsistent casing ('male', 'FEMALE', 'Fe-male'), leading whitespaces"],
        ["SeniorCitizen", "object", "Binary (0/1)", "Mixed types (integers 0/1, strings 'Yes'/'No', '0'/'1')"],
        ["TenureMonths", "float64", "Integer [0 - 72]", "Erroneous negative tenures (-5), impossible values (145), missing records"],
        ["MonthlyCharges", "object", "Float [$18 - $120]", "Currency prefixes ('$75.50'), negative values ('-45.0'), extreme outliers ($400+)"],
        ["TotalCharges", "object", "Float (Monetary)", "Whitespace empty strings (' '), missing values, currency symbols"],
        ["SignupDate", "object", "ISO Date (YYYY-MM-DD)", "Heterogeneous formats ('2021-04-12', '12/04/2021', 'Apr 12, 2021', 'INVALID_DATE')"],
        ["SatisfactionScore", "float64", "Integer [1 - 5]", "Out-of-bound sensor entries (0, 99), null entries"],
        ["Churn", "object", "Binary Flag", "Target noise ('Yes', 'No', 'yes', 'no', missing values)"]
    ]
    schema_table = doc.add_table(rows=1, cols=4)
    style_table(schema_table, [1.3, 0.9, 1.4, 2.9], schema_headers, schema_data)
    
    # Section 4: Data Quality Auditing & Missingness Analysis
    add_heading_styled(doc, "4. Exploratory Data Quality Auditing & Missingness Analysis", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "A critical pitfall in standard data cleaning workflows is evaluating null values solely via simplistic `.isnull().sum()` checks. "
        "Real-world data systems frequently contain hidden missingness—blank whitespace strings, sentinel placeholders ('N/A', 'NULL', 'INVALID_DATE'), "
        "and unparsed empty spaces. Our diagnostic audit uncovered that TotalCharges exhibited 113 explicit NaN values alongside 173 hidden whitespace entries, "
        "yielding a true missingness rate of 11.44%."
    )
    
    add_figure_with_caption(doc, "02_missing_values_bar.png", "Audit of True Missing & Incomplete Values Across Dataset Features")
    add_figure_with_caption(doc, "01_missing_data_matrix.png", "Binary Missingness Pattern Matrix Across 2,500 Ingested Observations")
    
    add_code_block(
        doc,
        "def audit_hidden_missingness(df):\n"
        "    missing_summary = []\n"
        "    for col in df.columns:\n"
        "        null_count = df[col].isnull().sum()\n"
        "        blank_count = 0\n"
        "        if df[col].dtype == 'object':\n"
        "            blank_count = df[col].apply(lambda x: 1 if isinstance(x, str) and \n"
        "                (x.strip() == '' or x.strip().upper() in ['NA', 'N/A', 'NULL', 'NONE']) else 0).sum()\n"
        "        total_missing = null_count + blank_count\n"
        "        missing_summary.append({\n"
        "            'Feature': col, 'Raw Nulls': null_count, 'Hidden Blanks': blank_count,\n"
        "            'Total Incomplete': total_missing, 'Missing %': round((total_missing / len(df)) * 100, 2)\n"
        "        })\n"
        "    return pd.DataFrame(missing_summary)",
        "Comprehensive Hidden Missingness and Sentinel Value Detector"
    )
    
    # Section 5: Systematic Data Cleaning & Remediation
    add_heading_styled(doc, "5. Systematic Data Cleaning & Value Remediation", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Data remediation was executed through a modular, deterministic strategy structured around five pillars:"
    )
    
    add_heading_styled(doc, "5.1 Primary Key Deduplication", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Customer records were validated for primary key uniqueness. Exactly 20 duplicate customer instances were identified and purged, "
        "retaining the primary first recorded entry to prevent artificial data leakage and inflated sample weighting."
    )
    
    add_heading_styled(doc, "5.2 Categorical Casing & Whitespace Standardization", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "All text attributes were subjected to bidirectional whitespace trimming (`.str.strip()`), case folding, and dictionary mapping. "
        "For example, the variable Gender contained seven distinct noisy representations ('Male', 'male', 'MALE', 'FEMALE', 'Female', 'female', 'Fe-male') "
        "which were normalized into a canonical binary domain {'Male', 'Female'}."
    )
    
    add_heading_styled(doc, "5.3 Numeric Parsing & Domain Constraint Enforcement", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Financial metrics formatted as strings with currency symbols and commas were parsed into 64-bit floating point representations. "
        "Furthermore, domain-specific physical constraints were strictly enforced:\n"
        "• Monthly Charges: Values < $0 (impossible negative billing) were flagged and converted to NaN for imputation.\n"
        "• Tenure Months: Values < 0 and values > 72 months (exceeding historical maximum operating window) were invalidated.\n"
        "• Satisfaction Ratings: Values outside the 1 to 5 Likert scale (e.g. 0, 99) were converted to NaN."
    )
    
    add_heading_styled(doc, "5.4 Multi-Format Datetime Normalization", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Signup dates received in non-standard mixed formats ('YYYY-MM-DD', 'DD/MM/YYYY', 'Month DD, YYYY') were parsed using flexible datetime parsing. "
        "Corrupted sentinel dates ('INVALID_DATE') were mathematically back-calculated based on the customer's recorded tenure duration relative to the anchor reference timestamp."
    )
    
    add_heading_styled(doc, "5.5 Advanced Imputation Strategy (KNN & Mode)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Rather than resorting to naive mean substitution (which artificially deflates feature variance and degrades correlation structures), "
        "we implemented K-Nearest Neighbors (KNN) Imputation (k=5, distance-weighted metric) across continuous numerical features "
        "(Tenure, Monthly Charges, Satisfaction Score). Categorical missing entries were filled using mode frequency."
    )
    
    add_code_block(
        doc,
        "# Advanced KNN Imputation for Multidimensional Numerical Features\n"
        "from sklearn.impute import KNNImputer\n\n"
        "num_cols = ['TenureMonths', 'MonthlyCharges', 'SatisfactionScore']\n"
        "knn_imputer = KNNImputer(n_neighbors=5, weights='distance')\n"
        "clean_df[num_cols] = knn_imputer.fit_transform(clean_df[num_cols])\n"
        "clean_df['SatisfactionScore'] = clean_df['SatisfactionScore'].round().astype(int)\n\n"
        "# Deterministic Logical Fill for Total Charges\n"
        "clean_df['TotalCharges'] = clean_df['TotalCharges'].fillna(clean_df['TenureMonths'] * clean_df['MonthlyCharges'])",
        "KNN Multi-Attribute Imputation Pipeline"
    )
    
    # Section 6: Outlier Diagnostics & Treatment
    add_heading_styled(doc, "6. Outlier Diagnostics, IQR Winsorization & Treatment", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Outliers can heavily bias parameter estimation in ordinary least squares (OLS) regression, distort gradient updates in neural networks, "
        "and skew centroid calculations in k-means clustering. We performed statistical outlier detection utilizing the Tukey 1.5 × IQR (Interquartile Range) method."
    )
    
    add_figure_with_caption(doc, "03_outliers_before_cleaning.png", "Boxplot Diagnostics of Numerical Features Prior to Outlier Remediation")
    
    # Outlier Stats Table
    outlier_headers = ["Metric / Feature", "Q1 (25th %)", "Q3 (75th %)", "IQR", "Lower Fence", "Upper Fence", "Outliers Capped"]
    outlier_data = [
        ["MonthlyCharges ($)", "$44.61", "$85.57", "$40.96", "$0.00", "$147.01", "12 records"],
        ["TotalCharges ($)", "$406.82", "$2,135.57", "$1,728.75", "$0.00", "$4,728.70", "148 records"]
    ]
    outlier_table = doc.add_table(rows=1, cols=7)
    style_table(outlier_table, [1.3, 0.8, 0.8, 0.8, 0.9, 0.9, 1.0], outlier_headers, outlier_data, align_right_cols=[1,2,3,4,5,6])
    
    p_treat = doc.add_paragraph()
    p_treat.add_run(
        "Rather than dropping outlier records—which would discard valid customer behavior data and induce survivor bias—we applied "
        "Winsorization (Robust Boundary Capping). Values exceeding the upper Tukey fence were clamped to the threshold, preserving the sample size of 2,480 valid customer profiles."
    )
    add_figure_with_caption(doc, "04_outlier_treatment_comparison.png", "Before vs. After Outlier Treatment Comparison for TotalCharges Feature")
    
    # Section 7: Skewness & Mathematical Transformations
    add_heading_styled(doc, "7. Distributional Skewness & Mathematical Transformations", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Highly skewed feature distributions violate the normality assumptions required by parametric algorithms and create severe instability in loss surfaces. "
        "The original TotalCharges feature exhibited a pronounced right-skewness of +1.063. We applied a natural logarithmic transform with unit shift (Log1p), "
        "transforming the distribution into a near-Gaussian profile and stabilizing the variance."
    )
    add_figure_with_caption(doc, "05_skewness_transformation.png", "Distribution & Kernel Density Estimation (KDE) Before and After Log1p Transformation")
    
    # Section 8: Feature Engineering & Preprocessing
    add_heading_styled(doc, "8. Feature Engineering, Categorical Encoding & Scaling", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "To maximize predictive signal for subsequent machine learning models, domain-specific feature engineering was conducted:"
    )
    
    eng_headers = ["Engineered Feature", "Mathematical / Logical Formulation", "Domain Rationale & Predictive Value"]
    eng_data = [
        ["TenureGroup", "Binned: [0-12m], [13-24m], [25-48m], [49-72m]", "Captures non-linear customer lifecycle stages (infant mortality vs. high brand loyalty)."],
        ["TotalServicesCount", "Sum of binary flags (Security + TechSupport + TV + Phone)", "Quantifies customer digital ecosystem stickiness; churn declines sharply with higher density."],
        ["MonthlyToTotalRatio", "MonthlyCharges / (TotalCharges + 1)", "Detects rapid billing shocks and newly acquired high-velocity accounts."],
        ["IsHighValueCustomer", "1 if MonthlyCharges > 75th percentile, else 0", "Enables targeted retention algorithms to prioritize high-margin churn risks."]
    ]
    eng_table = doc.add_table(rows=1, cols=3)
    style_table(eng_table, [1.5, 2.2, 2.8], eng_headers, eng_data)
    
    add_figure_with_caption(doc, "07_categorical_distributions.png", "Empirical Churn Probability Across Tenure Cohorts and Digital Service Densities")
    
    add_heading_styled(doc, "8.1 Categorical Encoding & Standardization", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "• Binary Categoricals (Gender, Partner, Dependents, PhoneService, PaperlessBilling) were mapped to integer flags {0, 1}.\n"
        "• Nominal Multi-Class Categoricals (Contract, InternetService, PaymentMethod, TenureGroup) were One-Hot Encoded with `drop_first=True` to prevent multicollinearity (dummy variable trap).\n"
        "• Continuous Features (Tenure, MonthlyCharges, TotalCharges_Log1p, CustomerAgeDays) were scaled using StandardScaler to zero mean and unit variance."
    )
    
    add_figure_with_caption(doc, "06_feature_correlation_matrix.png", "Full Correlation Matrix Across Preprocessed and Engineered Model-Ready Features")
    
    # Section 9: Challenges Faced & Overcome
    add_heading_styled(doc, "9. Challenges Faced & Solutions Implemented", level=1)
    
    add_callout_box(
        doc,
        "Challenge 1: Hidden Missingness in Total Charges",
        "Problem: TotalCharges contained whitespace characters (' ') which masqueraded as non-null strings, bypassing standard `.isna()` filters.\n"
        "Root Cause: Legacy data export systems emitted blank spaces when customer tenure was 0 months.\n"
        "Solution: Implemented a multi-tier regex cleaner converting whitespace patterns to NaN, followed by logical formula reconstruction (`TenureMonths * MonthlyCharges`).",
        "challenge"
    )
    
    add_callout_box(
        doc,
        "Challenge 2: Multi-Format Datetime Heterogeneity & Sentinel Dates",
        "Problem: Date strings arrived across four conflicting date formats alongside corrupted string markers ('INVALID_DATE').\n"
        "Solution: Constructed a flexible parser with format fallbacks. Corrupted dates were dynamically reconstructed using the reference cutoff date minus the customer's recorded tenure duration.",
        "challenge"
    )
    
    add_callout_box(
        doc,
        "Challenge 3: Preserving Sample Representation During Outlier Handling",
        "Problem: Raw truncation / row dropping of extreme financial charges would have removed over 6% of the customer base, discarding valuable high-revenue customer profiles.\n"
        "Solution: Replaced row deletion with Tukey 1.5 × IQR Winsorization and Log1p transformation, stabilizing feature scale while preserving 100% of validated customer observations.",
        "success"
    )
    
    # Section 10: Downstream Impact Analysis
    add_heading_styled(doc, "10. Downstream Impact on Machine Learning & Analytics", level=1)
    
    impact_headers = ["Algorithm Family", "Preprocessed Enhancements", "Downstream Performance Impact"]
    impact_data = [
        ["Linear & Logistic Models", "Multicollinearity removal, Log1p transform, Z-score scaling", "Prevents ill-conditioned Hessian matrices, eliminates coefficient inflation, improves gradient descent convergence rate."],
        ["Tree-Based Ensembles (XGBoost, RF)", "Categorical unification, high-value feature engineering", "Provides direct split nodes on engineered service density and tenure cohorts, reducing tree depth requirements."],
        ["Distance-Based (KNN, SVM, K-Means)", "StandardScaler normalization across continuous metrics", "Eliminates metric dominance where large dollar values ($4,000+) overpower small tenure intervals (1-72 months)."],
        ["Neural Networks / Deep Learning", "Zero mean / unit variance scaling, clean one-hot matrices", "Prevents vanishing / exploding gradients and stabilizes batch normalization layers."]
    ]
    impact_table = doc.add_table(rows=1, cols=3)
    style_table(impact_table, [1.6, 2.2, 2.7], impact_headers, impact_data)
    
    # Section 11: Conclusion & Deliverables
    add_heading_styled(doc, "11. Conclusion & Summary of Deliverables", level=1)
    p_concl = doc.add_paragraph()
    p_concl.add_run(
        "The Week 1 Data Acquisition, Cleaning, and Preprocessing task has been executed with industry-grade rigor. "
        "All data quality defects—including structural missingness, duplicate entities, out-of-range sensor/financial values, "
        "heterogeneous datetimes, and distribution outliers—have been systematically diagnosed, resolved, and validated. "
        "The resulting pipeline produces a robust 26-feature dataset ready for advanced predictive modeling and unsupervised customer segmentation."
    )
    
    add_callout_box(
        doc,
        "Deliverables Inventory",
        "1. Complete Technical Report: 'Week_1_Data_Acquisition_Cleaning_Report.docx'\n"
        "2. End-to-End Execution Script: 'data_preprocessing_pipeline.py'\n"
        "3. Interactive Notebook: 'Week_1_Data_Cleaning_and_Preprocessing.ipynb'\n"
        "4. Raw Dataset: 'raw_customer_dataset.csv' (2,500 rows, 20 features)\n"
        "5. Cleaned Dataset: 'cleaned_customer_dataset.csv' (2,480 rows, 22 features)\n"
        "6. ML-Ready Dataset: 'preprocessed_model_ready_dataset.csv' (2,480 rows, 26 features)\n"
        "7. Diagnostic Asset Visualizations: 8 High-Resolution PNGs in 'assets/' directory.",
        "success"
    )
    
    doc.save(DOC_PATH)
    print(f"[OK] Successfully built and saved publication-grade report to '{DOC_PATH}'.")

if __name__ == "__main__":
    build_report()
