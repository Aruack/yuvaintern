import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Sets background shading for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_callout_box(doc, text, title="STRATEGIC INSIGHT", bg_color="F0F9FF", border_color="0284C7"):
    """Creates a callout box with a styled border and background."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'<w:top w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run_t = p.add_run(f"📌 {title}\n")
    run_t.font.name = 'Calibri'
    run_t.font.size = Pt(10)
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(int(border_color[:2], 16), int(border_color[2:4], 16), int(border_color[4:], 16))
    
    run_body = p.add_run(text)
    run_body.font.name = 'Calibri'
    run_body.font.size = Pt(9.5)
    run_body.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    # Add spacing after table
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

def style_table(table, header_bg="1E293B", zebra_bg="F8FAFC", border_color="CBD5E1"):
    """Styles a python-docx table with enterprise headers, zebra stripes, and clean borders."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Set headers
    for cell in table.rows[0].cells:
        set_cell_background(cell, header_bg)
        set_cell_margins(cell, top=140, bottom=140, left=140, right=140)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                
    # Zebra striping for data rows
    for r_idx, row in enumerate(table.rows[1:], start=1):
        bg = zebra_bg if r_idx % 2 == 1 else "FFFFFF"
        for cell in row.cells:
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

def generate_document():
    doc = Document()
    
    # Page setup - Normal 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Color Palette: Deep Navy (#0F172A), Teal Accent (#0D9488), Dark Slate (#334155), Indigo (#4338CA)
    
    # -------------------------------------------------------------
    # TITLE & HEADER COVER BLOCK
    # -------------------------------------------------------------
    p_pre = doc.add_paragraph()
    p_pre.paragraph_format.space_before = Pt(0)
    p_pre.paragraph_format.space_after = Pt(2)
    run_pre = p_pre.add_run("VIRTUAL DATA SCIENCE EXPLORER INTERNSHIP | WEEK 1 DELIVERABLE")
    run_pre.font.name = 'Calibri'
    run_pre.font.size = Pt(9.5)
    run_pre.font.bold = True
    run_pre.font.color.rgb = RGBColor(0x0D, 0x94, 0x88) # Teal
    
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("Data Science Project Planning and Strategy Design")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Deep Navy
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(14)
    run_sub = p_sub.add_run("Project RetainAI: Enterprise Predictive Churn & Retention Optimization Engine")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    
    # Metadata summary table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Internship Program:", "Virtual Data Science Explorer Intern (YuvaIntern)"),
        ("Module & Milestone:", "Week 1: Project Conceptualization, Architecture & Strategy Design"),
        ("Primary Technology Stack:", "Python 3.11+, Scikit-Learn, XGBoost, LightGBM, MLflow, FastAPI, Streamlit"),
        ("Project Status & Allocation:", "Approved Strategy & Work Breakdown Structure (Total: 33.0 Allocated Hours)")
    ]
    for idx, (k, v) in enumerate(meta_data):
        cell_k = meta_table.cell(idx, 0)
        cell_v = meta_table.cell(idx, 1)
        cell_k.width = Inches(2.2)
        cell_v.width = Inches(4.3)
        
        p_k = cell_k.paragraphs[0]
        p_k.paragraph_format.space_before = Pt(2)
        p_k.paragraph_format.space_after = Pt(2)
        r_k = p_k.add_run(k)
        r_k.font.name = 'Calibri'
        r_k.font.bold = True
        r_k.font.size = Pt(9.5)
        r_k.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
        p_v = cell_v.paragraphs[0]
        p_v.paragraph_format.space_before = Pt(2)
        p_v.paragraph_format.space_after = Pt(2)
        r_v = p_v.add_run(v)
        r_v.font.name = 'Calibri'
        r_v.font.size = Pt(9.5)
        r_v.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        
        set_cell_background(cell_k, "F1F5F9")
        set_cell_background(cell_v, "FFFFFF")
        set_cell_margins(cell_k, 60, 60, 100, 100)
        set_cell_margins(cell_v, 60, 60, 100, 100)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECTION 1: INTRODUCTION & BACKGROUND
    # -------------------------------------------------------------
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("1. Introduction & Project Background")
    r_h1.font.name = 'Calibri'
    r_h1.font.size = Pt(16)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "In modern subscription-based business models (SaaS, FinTech, Telecommunications, and Digital Media), customer retention represents the single most critical driver of long-term unit economics and enterprise valuation. Acquiring a new customer is empirically estimated to cost between 5 to 7 times more than retaining an existing subscriber. However, conventional customer relationship management (CRM) workflows remain predominantly reactive: retention teams typically attempt rescue interventions only after a customer initiates cancellation or when usage drops to zero, at which point the probability of salvage is below 12%."
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "This project proposes "
    )
    r_bold = p.add_run("Project RetainAI")
    r_bold.font.bold = True
    p.add_run(
        ", a comprehensive, end-to-end predictive machine learning and customer retention optimization system designed in Python. By synthesizing multi-source telemetry data—including product engagement clickstream, subscription billing history, support ticket sentiment, and customer demographics—the platform forecasts subscriber churn risk 30 to 60 days in advance of potential attrition."
    )

    add_callout_box(
        doc,
        "By shifting customer success operations from reactive salvage campaigns to algorithmic, proactive early-warning triggers, enterprise organizations can reduce annual subscriber attrition by 18% to 25%, while optimizing promotional discount allocation through explainable AI (SHAP-driven prescriptive recommendations).",
        title="CORE BUSINESS JUSTIFICATION & VALUE HYPOTHESIS",
        bg_color="F0FDF4",
        border_color="16A34A"
    )

    # -------------------------------------------------------------
    # SECTION 2: PROJECT OBJECTIVES & SCOPE
    # -------------------------------------------------------------
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("2. Project Objectives & Scope Specification")
    r_h1.font.name = 'Calibri'
    r_h1.font.size = Pt(16)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run("To ensure rigorous execution and measurable business impact, the project establishes unambiguous SMART (Specific, Measurable, Achievable, Relevant, and Time-bound) objectives alongside clearly defined scope boundaries.")

    # Objectives Subheading
    h2 = doc.add_heading(level=2)
    r_h2 = h2.add_run("2.1 SMART Project Objectives")
    r_h2.font.name = 'Calibri'
    r_h2.font.size = Pt(13)
    r_h2.font.bold = True
    r_h2.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)

    obj_items = [
        ("Technical Model Performance:", "Develop and benchmark gradient-boosted classification models (XGBoost, LightGBM, CatBoost) achieving an Area Under the Precision-Recall Curve (PR-AUC) >= 0.82 and ROC-AUC >= 0.88 on an out-of-time test dataset, outperforming baseline logistic regression by at least 15%."),
        ("Lead Time & Proactive Window:", "Deliver accurate churn likelihood scores with an actionable prediction lead horizon of 30 to 60 days prior to contract renewal or anticipated drop-off date, allowing customer success representatives ample intervention time."),
        ("Model Interpretability & Prescriptive Rules:", "Integrate SHAP (SHapley Additive exPlanations) and LIME to generate granular, subscriber-level feature attribution charts, explaining top negative drivers (e.g., login inactivity, open support tickets, billing discrepancies)."),
        ("Operational Serving Architecture:", "Architect a low-latency Python FastAPI microservice capable of returning batch and real-time churn probabilities in under 200 milliseconds, coupled with an interactive Streamlit operations dashboard."),
        ("Financial Impact & ROI:", "Deliver an estimated net annual revenue preservation of $450,000 to $850,000 for a 100,000-subscriber cohort by prioritizing targeted retention incentives based on Customer Lifetime Value (CLV) weighting.")
    ]

    for title, desc in obj_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        r_t = p.add_run(title + " ")
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        p.add_run(desc)

    # In Scope vs Out of Scope Table
    h2 = doc.add_heading(level=2)
    r_h2 = h2.add_run("2.2 In-Scope vs. Out-of-Scope Boundary Matrix")
    r_h2.font.name = 'Calibri'
    r_h2.font.size = Pt(13)
    r_h2.font.bold = True
    r_h2.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)

    scope_table = doc.add_table(rows=6, cols=3)
    headers = ["Project Dimension", "In-Scope (Fully Addressed)", "Out-of-Scope (Future Iterations)"]
    for i, h in enumerate(headers):
        scope_table.cell(0, i).paragraphs[0].text = h
        
    scope_data = [
        ("Data Pipeline & Sources", "Ingestion of tabular structured CRM, billing history, aggregated app usage logs, and NLP sentiment scores from support ticket text.", "Real-time streaming raw clickstream ingestion via Apache Flink or distributed Kafka clusters at petabyte scale."),
        ("Feature Engineering", "Temporal aggregation (7-day, 30-day, 90-day velocity metrics), RFM scoring, ratio features, and interaction terms.", "Computer vision models or unstructured audio call recording transcription."),
        ("Algorithm Families", "Tree-based ensembles (XGBoost, LightGBM, CatBoost), regularized logistic baselines, random forests, and survival analysis.", "Deep multi-layer recurrent neural networks (RNN/LSTM) requiring dedicated multi-GPU cluster hardware."),
        ("Deployment & Serving", "Containerized FastAPI REST API endpoints, Docker containerization, Streamlit visualization UI, and SQLite/PostgreSQL database.", "Multi-region Kubernetes auto-scaling orchestration and enterprise SSO/LDAP corporate authentication integration."),
        ("Model Monitoring", "Evidently AI automated drift detection reports (covariate shift, PSI, KS-test) and automated retraining threshold triggers.", "Automated zero-human self-healing model updates directly to live production financial systems.")
    ]

    for row_idx, data in enumerate(scope_data, start=1):
        for col_idx, text in enumerate(data):
            scope_table.cell(row_idx, col_idx).paragraphs[0].text = text

    style_table(scope_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # -------------------------------------------------------------
    # SECTION 3: METHODOLOGY & STRATEGY
    # -------------------------------------------------------------
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("3. Methodology & Strategic Architecture")
    r_h1.font.name = 'Calibri'
    r_h1.font.size = Pt(16)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "The project adopts an extended CRISP-DM (Cross-Industry Standard Process for Data Mining) methodology, augmented with modern MLOps principles, automated data validation, and feedback monitoring loops. Below is the end-to-end data science lifecycle and pipeline architecture developed for Project RetainAI."
    )

    # Embed Figure 1
    if os.path.exists("figures/figure1_lifecycle_flowchart.png"):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(2)
        doc.add_picture("figures/figure1_lifecycle_flowchart.png", width=Inches(6.5))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        r_cap = p_cap.add_run("Figure 1: Project RetainAI End-to-End Data Science Lifecycle & Pipeline Flowchart")
        r_cap.font.name = 'Calibri'
        r_cap.font.size = Pt(9)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Detailed Methodology Phases
    phases = [
        ("Phase 1: Data Acquisition, Ingestion & Schema Contracts",
         "A rigorous data collection strategy combining three fundamental data silos:\n"
         "1. Customer Profile & Demographics (Account tenure, contract type, payment method, geographic territory, industry vertical).\n"
         "2. Behavioral Engagement Telemetry (Daily active minutes, feature adoption breadth, login frequency, API call velocity over 7d/30d/90d intervals).\n"
         "3. Customer Service & Support Signals (Ticket volume, resolution time, customer satisfaction scores (CSAT), and NLP sentiment polarity derived from ticket summaries).\n"
         "Data integrity is established using Great Expectations to validate schema types, non-null constraints, and valid ranges before ingestion into staging tables."),
        
        ("Phase 2: Data Preprocessing, Cleaning & Pipeline Design",
         "A modular, leakage-free data cleaning pipeline built using Scikit-Learn Pipelines and Polars/Pandas:\n"
         "• Missing Data Imputation: Categorical variables imputed with dedicated 'Missing' tokens; numerical values imputed using median/iterative KNN based on subscriber tier cohorts.\n"
         "• Outlier Treatment: Winsorization (capping at 1st and 99th percentiles) on skewed numerical metrics (e.g., data consumption, session length) to stabilize gradient descent without discarding critical power-user data.\n"
         "• Leakage Prevention: Strict temporal splitting (Training on Months 1-8, Validation on Months 9-10, Testing on Months 11-12) to ensure no future behavioral information leaks into historical training matrices."),
        
        ("Phase 3: Exploratory Data Analysis (EDA) & Behavioral Profiling",
         "Systematic statistical exploration to uncover churn dynamics:\n"
         "• Univariate & Bivariate Distributions: Identifying distribution skewness, high-kurtosis engagement metrics, and class distribution (e.g., 16% baseline churn rate).\n"
         "• Cohort Retention Heatmaps: Analyzing subscriber retention curves grouped by acquisition channel and onboarding vintage.\n"
         "• Correlation & Collinearity Screening: Calculating Pearson/Spearman correlation matrices and Variance Inflation Factors (VIF) to eliminate redundant collinear predictors."),
        
        ("Phase 4: Advanced Feature Engineering & Feature Store Formulation",
         "Transforming raw transactional records into high-signal predictive features:\n"
         "• Recency, Frequency, Monetary (RFM) Scores: Quantifying recent activity recency (days since last login), transaction frequency, and total lifetime spend.\n"
         "• Velocity & Momentum Indicators: Computing rolling ratio features (e.g., [Activity Last 7 Days] / [Activity Last 30 Days]) to detect sudden drop-offs in product usage.\n"
         "• Contract & Payment Health: Days until renewal, count of failed billing retries, and upgrade/downgrade event counts.\n"
         "• Text NLP Features: DistilBERT / VADER sentiment polarity and complaint keyword counts extracted from customer service transcripts."),
        
        ("Phase 5: Machine Learning Modeling Strategy & Hyperparameter Optimization",
         "A multi-model comparative architecture:\n"
         "• Baseline Model: Regularized Logistic Regression with L1/L2 penalties and standard scaling to establish an interpretable benchmark.\n"
         "• Primary Classifiers: Extreme Gradient Boosting (XGBoost), LightGBM, and CatBoost (utilizing native categorical handling).\n"
         "• Survival Analysis: Cox Proportional Hazards model to estimate time-to-churn and dynamic hazard rates across subscriber lifespans.\n"
         "• Optimization: Bayesian Hyperparameter Optimization via Optuna (50 trials) optimizing PR-AUC under 5-fold Stratified Time-Series Split cross-validation."),
        
        ("Phase 6: Model Evaluation, Validation & Business Cost Matrix",
         "Evaluation tailored to imbalanced classification and financial business impact:\n"
         "• Primary Technical Metrics: Precision-Recall AUC (PR-AUC), Receiver Operating Characteristic AUC (ROC-AUC), F1-Score at optimal classification threshold, and Brier Reliability Score.\n"
         "• Business Cost-Utility Matrix: Explicitly evaluating False Negatives (unidentified churner losing full customer lifetime value) versus False Positives (unnecessary promotional discount spend).\n"
         "• Cumulative Gains & Decile Lift Charts: Quantifying the proportion of total churners captured within the top 20% of highest predicted risk deciles.")
    ]

    for title, body in phases:
        h2 = doc.add_heading(level=2)
        r_h2 = h2.add_run(title)
        r_h2.font.name = 'Calibri'
        r_h2.font.size = Pt(12)
        r_h2.font.bold = True
        r_h2.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(body)

    # -------------------------------------------------------------
    # SECTION 4: SYSTEM ARCHITECTURE & SERVING BLUEPRINT
    # -------------------------------------------------------------
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("4. Production System Architecture & Serving Strategy")
    r_h1.font.name = 'Calibri'
    r_h1.font.size = Pt(16)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "Deploying the predictive churn engine requires a multi-tier production architecture balancing low-latency real-time inference with robust batch scoring pipelines, experiment tracking, and continuous data drift observability."
    )

    # Embed Figure 3
    if os.path.exists("figures/figure3_system_architecture.png"):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(2)
        doc.add_picture("figures/figure3_system_architecture.png", width=Inches(6.5))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        r_cap = p_cap.add_run("Figure 2: Project RetainAI Multi-Tier Production System & Serving Blueprint")
        r_cap.font.name = 'Calibri'
        r_cap.font.size = Pt(9)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    arch_points = [
        ("Inference Microservice (FastAPI):", "Lightweight, asynchronous REST API containerized with Docker. Endpoints include `/predict_single` for on-demand customer evaluation and `/batch_score` for weekly bulk account scoring."),
        ("Explainability Engine (SHAP):", "Generates real-time force plots and feature contribution vectors accompanying each churn probability output, allowing account managers to immediately see why a user is at risk."),
        ("Operational Dashboard (Streamlit):", "Interactive web UI providing customer success teams with filtered account risk rankings, revenue-at-risk decile summaries, and recommended intervention playbooks."),
        ("Continuous Observability (Evidently AI):", "Automated weekly monitoring tracking Population Stability Index (PSI) and Kolmogorov-Smirnov (KS) tests across key input features to detect covariate and concept drift.")
    ]

    for title, desc in arch_points:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        r_t = p.add_run(title + " ")
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        p.add_run(desc)

    # -------------------------------------------------------------
    # SECTION 5: TIMELINE & TOOLS ECOSYSTEM
    # -------------------------------------------------------------
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("5. Work Breakdown Structure, Timeline & Toolchain")
    r_h1.font.name = 'Calibri'
    r_h1.font.size = Pt(16)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "To satisfy the 30-35 hour project effort allocation, the implementation plan is divided into five structured phases executed sequentially over a two-week sprint schedule (averaging 3.3 hours/day or dedicated intensive blocks)."
    )

    # Embed Figure 2 Gantt Chart
    if os.path.exists("figures/figure2_project_gantt_timeline.png"):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(2)
        doc.add_picture("figures/figure2_project_gantt_timeline.png", width=Inches(6.5))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        r_cap = p_cap.add_run("Figure 3: Project RetainAI 30-35 Hour Work Breakdown Structure & Timeline Gantt Chart")
        r_cap.font.name = 'Calibri'
        r_cap.font.size = Pt(9)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Detailed Timeline Table
    timeline_table = doc.add_table(rows=6, cols=4)
    tl_headers = ["Phase & Focus Area", "Hours Allocated", "Key Activities & Work Tasks", "Core Milestone Deliverable"]
    for i, h in enumerate(tl_headers):
        timeline_table.cell(0, i).paragraphs[0].text = h

    tl_data = [
        ("Phase 1: Project Scoping & Data Architecture Specs", "5.0 Hours\n(Hours 0 - 5)", 
         "• Define business problem & KPI hierarchy\n• Establish data dictionaries & schema contracts\n• Configure Git repository & environment dependencies", 
         "Milestone 1: Project Charter & Data Architecture Spec Approved"),
        ("Phase 2: Data Preprocessing, Cleaning & Pipeline Design", "6.5 Hours\n(Hours 5 - 11.5)", 
         "• Implement missing data & outlier handling modules\n• Build automated Great Expectations validation suite\n• Formulate temporal train/val/test data splitting logic", 
         "Milestone 2: Modular Clean Data Pipeline Verified & Tested"),
        ("Phase 3: Exploratory Data Analysis & Feature Store Strategy", "7.0 Hours\n(Hours 11.5 - 18.5)", 
         "• Conduct univariate, bivariate, & cohort retention EDA\n• Engineer RFM, engagement velocity, & sentiment features\n• Multi-collinearity screening and feature selection", 
         "Milestone 3: Engineered Feature Store & Baseline Benchmark"),
        ("Phase 4: ML Modeling, Hyperparameter Tuning & Evaluation", "8.0 Hours\n(Hours 18.5 - 26.5)", 
         "• Train XGBoost, LightGBM, CatBoost, & Logistic models\n• Run Optuna Bayesian hyperparameter search (50 trials)\n• Generate PR-AUC, ROC-AUC, Lift curves & Cost Matrix", 
         "Milestone 4: Validated Champion Model with Superior Lift"),
        ("Phase 5: MLOps Architecture, Explainability & Deliverables", "6.5 Hours\n(Hours 26.5 - 33.0)", 
         "• Build FastAPI REST service & Streamlit analytics UI\n• Integrate SHAP waterfall explainability charts\n• Draft executive strategy document & final code freeze", 
         "Milestone 5: Production-Ready Strategy Document & Demo")
    ]

    for row_idx, data in enumerate(tl_data, start=1):
        for col_idx, text in enumerate(data):
            timeline_table.cell(row_idx, col_idx).paragraphs[0].text = text

    style_table(timeline_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Python Technology Stack Matrix
    h2 = doc.add_heading(level=2)
    r_h2 = h2.add_run("5.2 Comprehensive Python Ecosystem & Technology Stack")
    r_h2.font.name = 'Calibri'
    r_h2.font.size = Pt(13)
    r_h2.font.bold = True
    r_h2.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)

    tools_table = doc.add_table(rows=7, cols=3)
    t_headers = ["Functional Layer", "Selected Python Libraries / Tools", "Strategic Rationale & Role"]
    for i, h in enumerate(t_headers):
        tools_table.cell(0, i).paragraphs[0].text = h

    t_data = [
        ("Data Manipulation & Ingestion", "Pandas, Polars, SQLAlchemy, PyArrow", "High-throughput data parsing, memory-efficient columnar operations, and structured relational database connectivity."),
        ("Data Quality & Testing", "Great Expectations, PyTest, Pydantic", "Automated data contract validation, schema enforcement, and test-driven regression suites for data pipelines."),
        ("EDA & Statistical Visualization", "Matplotlib, Seaborn, Plotly, SciPy", "Publication-ready distribution plots, correlation matrices, interactive drill-downs, and hypothesis testing."),
        ("Feature Engineering & Machine Learning", "Scikit-Learn, XGBoost, LightGBM, CatBoost, Lifelines", "Modular transformer pipelines, state-of-the-art gradient boosting algorithms, and survival analysis models."),
        ("Hyperparameter Optimization", "Optuna, Hyperopt", "Automated Bayesian optimization with efficient tree-structured Parzen estimators and early stopping pruning."),
        ("Interpretability, Serving & MLOps", "SHAP, LIME, FastAPI, Uvicorn, Streamlit, MLflow, Evidently AI", "Local and global feature attribution, asynchronous REST inference, operational dashboarding, experiment tracking, and drift detection.")
    ]

    for row_idx, data in enumerate(t_data, start=1):
        for col_idx, text in enumerate(data):
            tools_table.cell(row_idx, col_idx).paragraphs[0].text = text

    style_table(tools_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # -------------------------------------------------------------
    # SECTION 6: EXPECTED OUTCOMES & RISK MANAGEMENT
    # -------------------------------------------------------------
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("6. Expected Outcomes, Key Metrics & Risk Management")
    r_h1.font.name = 'Calibri'
    r_h1.font.size = Pt(16)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "A robust data science plan must proactively quantify expected business impact and systematically mitigate technical and operational risks."
    )

    # Embed Figure 4 Risk Matrix
    if os.path.exists("figures/figure4_risk_matrix.png"):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(2)
        doc.add_picture("figures/figure4_risk_matrix.png", width=Inches(5.5))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        r_cap = p_cap.add_run("Figure 4: Strategic Risk Assessment & Mitigation Framework")
        r_cap.font.name = 'Calibri'
        r_cap.font.size = Pt(9)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Risk Mitigation Table
    risk_table = doc.add_table(rows=7, cols=4)
    r_headers = ["Risk ID & Category", "Identified Risk Event", "Severity & Likelihood", "Concrete Mitigation Strategy"]
    for i, h in enumerate(r_headers):
        risk_table.cell(0, i).paragraphs[0].text = h

    r_data = [
        ("R1: Data Leakage", "Future engagement signals or target leakage across temporal windows in train/test splits.", "High Impact\nLow Likelihood", 
         "Enforce strict out-of-time (OOT) temporal splits; compute rolling window aggregates strictly up to prediction cut-off date (T_0)."),
        ("R2: Class Imbalance", "Low positive churn event prevalence (e.g. 10-15%) distorting default accuracy metrics.", "High Impact\nMedium Likelihood", 
         "Optimize for PR-AUC and Cost-Weighted F1 rather than accuracy; apply scale_pos_weight parameter in XGBoost/LightGBM and SMOTE-Tomek."),
        ("R3: Production Drift", "Consumer behavior shifts post-marketing campaign degrading model predictive accuracy over time.", "Medium Impact\nMedium Likelihood", 
         "Deploy Evidently AI to monitor Population Stability Index (PSI > 0.2 threshold) and configure automated alerts for weekly retraining jobs."),
        ("R4: Inference Latency", "Complex feature pipelines causing response delays (>500ms) in real-time API endpoints.", "Low Impact\nLow Likelihood", 
         "Pre-compute static daily features in a feature store (Feast/Redis); optimize FastAPI payloads with asynchronous workers."),
        ("R5: Retention Bias", "Algorithms disproportionately offering retention discounts to specific demographic tiers.", "Medium Impact\nLow Likelihood", 
         "Conduct disparate impact audits across sensitive demographic features using Fairlearn to ensure equitable intervention eligibility."),
        ("R6: User Adoption", "Customer success representatives reluctant to trust complex black-box machine learning predictions.", "Low Impact\nLow Likelihood", 
         "Integrate transparent SHAP waterfall reason codes in the Streamlit UI, explaining the specific actionable triggers behind every high-risk alert.")
    ]

    for row_idx, data in enumerate(r_data, start=1):
        for col_idx, text in enumerate(data):
            risk_table.cell(row_idx, col_idx).paragraphs[0].text = text

    style_table(risk_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Expected Outcomes KPI summary
    h2 = doc.add_heading(level=2)
    r_h2 = h2.add_run("6.2 Target Success Metrics & Business ROI")
    r_h2.font.name = 'Calibri'
    r_h2.font.size = Pt(13)
    r_h2.font.bold = True
    r_h2.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)

    kpi_items = [
        ("Churn Rate Reduction:", "Target reduction in annual gross customer churn from baseline 18% to 14.5% (a 19.4% relative reduction)."),
        ("Intervention Lead Time:", "Achieve an average warning horizon of 45 days prior to contract termination or inactivity cut-off."),
        ("Top-Decile Churn Capture (Lift):", "Capture >= 55% of all churning accounts within the top 20% highest-risk predicted segment (2.75x lift over random selection)."),
        ("Net Saved Annual Recurring Revenue (ARR):", "Estimated net ARR preservation of $620,000 annually per 100,000 accounts after accounting for retention offer expenses.")
    ]

    for title, desc in kpi_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        r_t = p.add_run(title + " ")
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        p.add_run(desc)

    # -------------------------------------------------------------
    # SECTION 7: CONCLUSION & TRANSITION TO WEEK 2
    # -------------------------------------------------------------
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("7. Conclusion & Next Steps (Transitioning to Week 2)")
    r_h1.font.name = 'Calibri'
    r_h1.font.size = Pt(16)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "This Week 1 project plan and strategy design establishes an exhaustive, technically rigorous, and financially grounded foundation for Project RetainAI. By defining clear scope boundaries, architecting an end-to-end data pipeline, establishing an allocated 33-hour work breakdown structure, and anticipating deployment and drift risks, the project is fully positioned for seamless technical execution in Week 2 (Data Collection, Preprocessing & Exploratory Analysis)."
    )

    add_callout_box(
        doc,
        "With strategy, timeline, schemas, and architecture formally documented, the next immediate technical milestones for Week 2 will include initializing the synthetic multi-source data generation scripts, configuring the automated PyTest/Great Expectations data validation harnesses, and conducting baseline EDA notebooks.",
        title="EXECUTIVE SIGN-OFF & TRANSITION READINESS",
        bg_color="FDF4FF",
        border_color="9333EA"
    )

    # Save document
    output_filename = "Project_Planning_and_Strategy_Design.docx"
    doc.save(output_filename)
    print(f"Successfully generated DOCX file: {output_filename}")

if __name__ == "__main__":
    generate_document()
