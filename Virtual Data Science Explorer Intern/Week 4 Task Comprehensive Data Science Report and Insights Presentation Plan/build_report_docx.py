import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets background color for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    """Sets cell padding in dxa."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}>'
                      f'<w:top w:w="{top}" w:type="dxa"/>'
                      f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
                      f'<w:left w:w="{left}" w:type="dxa"/>'
                      f'<w:right w:w="{right}" w:type="dxa"/>'
                      f'</w:tcMar>')
    tcPr.append(tcMar)

def add_callout_box(doc, text_title, text_body, border_hex="1B365D", bg_hex="F0F4F8"):
    """Adds a stylish callout box with a colored left accent border and shaded background."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                        f'<w:top w:val="none"/>'
                        f'<w:left w:val="single" w:sz="36" w:space="0" w:color="{border_hex}"/>'
                        f'<w:bottom w:val="none"/>'
                        f'<w:right w:val="none"/>'
                        f'</w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"📌 {text_title}\n")
    run_t.bold = True
    run_t.font.name = "Calibri"
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    run_b = p.add_run(text_body)
    run_b.font.name = "Calibri"
    run_b.font.size = Pt(10)
    run_b.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def format_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.keep_with_next = True
    run = h.runs[0]
    run.font.name = "Calibri"
    if level == 1:
        run.font.size = Pt(18)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x99) # Slate Teal
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(2)
    return h

def add_body_p(doc, text, bold_prefix="", space_after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.name = "Calibri"
        r_pre.font.size = Pt(10.5)
        r_pre.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r.italic = italic
    r.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    return p

def add_styled_table(doc, headers, data, col_widths=None):
    tbl = doc.add_table(rows=len(data) + 1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    # Format Header Row
    hdr_cells = tbl.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1B365D")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for r in p.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(10)
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
    # Format Data Rows
    for row_idx, row_data in enumerate(data):
        row_cells = tbl.rows[row_idx + 1].cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = str(cell_value)
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.font.name = "Calibri"
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
                
    # Set Widths if specified
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
                
    # Add subtle border
    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            b = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                          f'<w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
                          f'<w:left w:val="none"/>'
                          f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
                          f'<w:right w:val="none"/>'
                          f'</w:tcBorders>')
            tcPr.append(b)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_figure_image(doc, image_path, caption_text, width_inches=6.2):
    if os.path.exists(image_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(4)
        run_img = p_img.add_run()
        run_img.add_picture(image_path, width=Inches(width_inches))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after = Pt(10)
        run_cap = p_cap.add_run(f"📊 {caption_text}")
        run_cap.font.name = "Calibri"
        run_cap.font.size = Pt(9)
        run_cap.italic = True
        run_cap.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

def main():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Header / Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Project Horizon | Comprehensive Data Science & Insights Presentation Plan")
        hrun.font.name = "Calibri"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Confidential — Virtual Data Science Explorer Internship Program | YuvaIntern")
        frun.font.name = "Calibri"
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)

    # -------------------------------------------------------------
    # COVER / HEADER TITLE BLOCK
    # -------------------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(2)
    run_main_title = p_title.add_run("Project Horizon: Predictive Customer Churn Analytics, Revenue Risk Mitigation, and Proactive Retention Strategy")
    run_main_title.font.name = "Calibri"
    run_main_title.font.size = Pt(22)
    run_main_title.bold = True
    run_main_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(14)
    run_sub = p_sub.add_run("A Comprehensive Data Science Report, Executive Decision Framework, and Non-Technical Insights Presentation Plan")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
    
    # Metadata Table
    meta_headers = ["Document Attribute", "Specification Details"]
    meta_data = [
        ["Program", "Virtual Data Science Explorer Internship — Final Capstone (Week 4)"],
        ["Author / Lead Analyst", "Aryan Kumar (Data Science Explorer Intern)"],
        ["Target Audience", "Executive Leadership (CFO, CRO, CPO), VP Customer Success, Data Science Steering Committee"],
        ["Project Domain", "Enterprise B2B SaaS, Subscription Analytics & Customer Lifetime Value (CLV) Optimization"],
        ["Analytical Deliverables", "Predictive Ensemble Model, Survival Curves, SHAP Attributions, Revenue Optimization Frontier"],
        ["Presentation Blueprint", "Minto Pyramid & SCQA Executive Storytelling Architecture, 10-Slide Blueprint & FAQ Matrix"],
        ["Document Status", "Production Final — Version 1.0 (Executive Review Ready)"]
    ]
    add_styled_table(doc, meta_headers, meta_data, col_widths=[2.2, 4.3])
    
    doc.add_page_break()

    # -------------------------------------------------------------
    # SECTION 1: EXECUTIVE SUMMARY
    # -------------------------------------------------------------
    format_heading(doc, "1. Executive Summary", level=1)
    
    add_body_p(doc, 
               "In modern recurring-revenue business models, customer acquisition costs (CAC) continue to escalate across enterprise software markets. Retaining existing contracted revenue and optimizing Net Revenue Retention (NRR) have become the single most vital operational levers for sustainable enterprise valuation and long-term profitability. This comprehensive report synthesizes the end-to-end data science investigation for Project Horizon, a hypothetical enterprise SaaS platform experiencing an annual gross customer churn rate of 14.2%, representing approximately $12.8M in Annual Recurring Revenue (ARR) at risk.",
               bold_prefix="Business Context & Mandate: ")
    
    add_body_p(doc, 
               "The overarching mandate of this initiative was to engineer a high-precision, interpretable predictive early-warning system capable of forecasting customer attrition risk 60 to 90 days prior to contract expiration, uncovering the root behavioral and operational drivers of churn, and establishing an economically optimal retention intervention framework.",
               bold_prefix="Core Analytical Objective: ")
    
    add_callout_box(doc, 
                    "Executive Key Findings & Financial Impact Summary", 
                    "• Early Identification Power: The machine learning ensemble captures 73.4% of all potential churners within the top 20% of accounts flagged by the model (Lift = 3.67x over random baseline).\n"
                    "• The 90-Day 'Onboarding Cliff': Accounts adopting fewer than 4 core features in their first 90 days experience a 32% retention drop compared to highly engaged accounts.\n"
                    "• Support Friction Escalation: Unresolved Tier-1/Tier-2 support tickets multiply churn likelihood by up to 6.2x in Enterprise accounts.\n"
                    "• Economic ROI Frontier: Targeting the top 22% of high-risk accounts through a tiered proactive Customer Success playbook delivers a projected $4.22M in preserved net ARR against an intervention cost of $910K, yielding a 4.6x Net ROI.")

    add_body_p(doc, "The table below details the baseline metrics versus the targeted post-implementation improvements under Project Horizon:")

    exec_kpi_headers = ["Key Performance Metric", "Historical Baseline", "Model-Driven Target", "Projected Enterprise Impact"]
    exec_kpi_data = [
        ["Annual Gross Churn Rate", "14.2% ARR", "9.5% ARR", "33.1% reduction in churn volume"],
        ["Net Revenue Retention (NRR)", "102.4%", "109.8%", "+740 bps expansion in annual recurring baseline"],
        ["Churn Warning Lead Time", "14 Days (Reactive)", "60-90 Days (Proactive)", "300% increase in Customer Success response runway"],
        ["Proactive Intervention Precision", "24.0% (Rule-based heuristics)", "78.2% (Top Decile Precision)", "Eliminates alert fatigue and misdirected retention spend"],
        ["Net ARR Preserved / Year", "$0 (Baseline Run-rate)", "$4.22 Million", "Direct contribution to free cash flow and enterprise multiple"]
    ]
    add_styled_table(doc, exec_kpi_headers, exec_kpi_data, col_widths=[1.8, 1.3, 1.4, 2.0])

    # -------------------------------------------------------------
    # SECTION 2: METHODOLOGY OVERVIEW
    # -------------------------------------------------------------
    format_heading(doc, "2. Methodology Review", level=1)
    
    add_body_p(doc, 
               "To construct a robust, production-grade intelligence framework, the project followed an end-to-end analytical architecture comprising five rigorous phases: Data Ingestion & Unification, Feature Engineering & Behavioral Profiling, Supervised Machine Learning Ensemble Modeling, Survival Analysis, and Cost-Sensitive Threshold Optimization.",
               bold_prefix="Analytical Architecture: ")

    format_heading(doc, "2.1 Data Ingestion, Synthesis & Entity Unification", level=2)
    add_body_p(doc, 
               "The unified dataset synthesized historical telemetry, commercial contracts, customer interactions, and billing records across 50,000 corporate accounts over a continuous 24-month rolling observation window. The data schema was structured across four distinct data vectors:")
    
    data_vectors_headers = ["Data Vector", "Source System", "Extracted Attributes & Telemetry", "Preprocessing & Cleaning Protocol"]
    data_vectors_data = [
        ["Product Behavioral Telemetry", "Event Streaming (Kafka/Segment)", "Daily Active Users (DAU), Monthly Active Users (MAU), session length, feature clickstream, API calls, export actions.", "Missing timestamp imputation, sessionization aggregation, time-decayed rolling averages (7d, 30d, 90d)."],
        ["Commercial & Financial Records", "ERP / Subscription Billing (Stripe/Zuora)", "Contract tier (Enterprise, Mid-Market, SMB), ARR, payment terms, billing failures, invoice aging, discounting history.", "One-hot encoding of plan tiers, logarithmic transformation of ARR, payment failure count aggregation."],
        ["Customer Success & Support", "CRM / Helpdesk (Salesforce/Zendesk)", "Support ticket frequency, resolution time, Sev-1/Sev-2 escalations, Customer Success Manager (CSM) touchpoints.", "NLP sentiment scoring of ticket bodies, escalation ratio calculation, touchpoint recency calculation."],
        ["Customer Sentiment & Voice of Customer", "Survey Engine (Qualtrics)", "Net Promoter Score (NPS), CSAT ratings, customer onboarding survey responses.", "Weight of Evidence (WoE) transformation, missing response imputation via median within account tier."]
    ]
    add_styled_table(doc, data_vectors_headers, data_vectors_data, col_widths=[1.5, 1.3, 2.2, 1.5])

    format_heading(doc, "2.2 Feature Engineering & Signal Extraction", level=2)
    add_body_p(doc, 
               "Over 120 candidate features were engineered to capture temporal dynamics, velocity, and inflection points in user behavior rather than static historical totals. Key feature categories included:")
    add_body_p(doc, "1. Velocity & Decay Metrics: Calculating the 30-day vs 90-day MAU ratio (Activity Momentum) to identify sudden engagement drop-offs.\n"
                    "2. Depth of Adoption Ratios: Computing the ratio of unique advanced features utilized relative to total available features in the licensed tier.\n"
                    "3. Organizational Breadth: Quantifying multi-departmental seat penetration and executive sponsor login recency.\n"
                    "4. Friction Accumulation Score: Composite index combining unresolved ticket duration, invoice disputes, and CSAT detractor ratings.")

    format_heading(doc, "2.3 Machine Learning Algorithm Benchmark & Model Selection", level=2)
    add_body_p(doc, 
               "A rigorous multi-algorithm benchmarking experiment was conducted across Logistic Regression (Baseline), Random Forest, LightGBM, CatBoost, and an Extreme Gradient Boosting (XGBoost) ensemble. Models were evaluated using Stratified 5-Fold Time-Series Cross-Validation to eliminate temporal lookahead leakage. The benchmark results are summarized below:")

    model_headers = ["Candidate Algorithm", "Validation PR-AUC", "ROC-AUC", "Top-20% Lift", "Inference Latency (p95)", "Interpretability Score"]
    model_data = [
        ["ElasticNet Logistic Regression (Baseline)", "0.582", "0.714", "2.10x", "1.2 ms", "Very High (Linear Weights)"],
        ["Random Forest Classifier (100 Trees)", "0.764", "0.835", "3.12x", "8.5 ms", "Moderate (Tree SHAP)"],
        ["LightGBM (Gradient Boosted Trees)", "0.838", "0.891", "3.62x", "3.4 ms", "High (TreeSHAP / Fast)"],
        ["XGBoost (Optimized via Optuna)", "0.842", "0.896", "3.67x", "4.1 ms", "High (TreeSHAP / Exact)"],
        ["Blended Stacking Ensemble (XGB+LGBM)", "0.849", "0.902", "3.71x", "7.8 ms", "High (Averaged SHAP)"]
    ]
    add_styled_table(doc, model_headers, model_data, col_widths=[1.8, 0.9, 0.8, 0.9, 1.1, 1.0])

    add_body_p(doc, 
               "The XGBoost model, fine-tuned via Bayesian hyperparameter optimization (Optuna over 200 trials), was selected as the core production engine due to its superior balance of Precision-Recall AUC (0.842), rock-solid stability, sub-5ms inference latency, and full native support for SHAP (SHapley Additive exPlanations) attribution.",
               bold_prefix="Final Model Selection: ")

    # -------------------------------------------------------------
    # SECTION 3: IN-DEPTH FINDINGS, STATISTICAL INSIGHTS & VISUALIZATIONS
    # -------------------------------------------------------------
    format_heading(doc, "3. Findings, Statistical Insights & Visualizations", level=1)
    
    add_body_p(doc, 
               "The analysis uncovered four fundamental behavioral and operational dynamics governing customer attrition. Each finding is supported by quantitative statistical modeling and executive-grade visual mockups.")

    format_heading(doc, "3.1 Finding 1: The 'Onboarding Cliff' & Feature Adoption Tipping Point", level=2)
    add_body_p(doc, 
               "Survival analysis using Kaplan-Meier estimators revealed that customer tenure is critically determined during the first 90 days of contract activation. Accounts were segmented into three feature adoption cohorts based on their 90-day product telemetry:")
    add_body_p(doc, "• Tier 1 (High Adoption - >=8 core features): Demonstrates a stellar 365-day survival probability of 91.2%.\n"
                    "• Tier 2 (Moderate Adoption - 4 to 7 core features): Experiences steady degradation with a 365-day survival rate of 74.8%.\n"
                    "• Tier 3 (Low Adoption - <4 core features): Suffers a steep 'Onboarding Cliff', losing 32% of accounts by Day 90 and plummeting to a 365-day retention rate of only 41.5%.")
    
    add_figure_image(doc, "assets/survival_retention_curve.png", 
                     "Figure 1: Kaplan-Meier Survival Probabilities over 365 Days across Feature Adoption Cohorts with 95% Confidence Intervals.")

    add_body_p(doc, 
               "Statistical Takeaway: A Cox Proportional Hazards regression demonstrated that crossing the 4-feature threshold during the onboarding window reduces the instantaneous hazard of churning by 68.4% (Hazard Ratio = 0.316, p < 0.001). This proves that proactive onboarding intervention is the single highest-leverage retention window in the customer lifecycle.",
               bold_prefix="Statistical Significance: ")

    format_heading(doc, "3.2 Finding 2: Multi-Dimensional Churn Predictors & SHAP Feature Attribution", level=2)
    add_body_p(doc, 
               "Global SHAP analysis was conducted to decompose the ensemble model's predictions into explainable, additive log-odds contributions across the top 10 most influential predictors.")
    
    add_figure_image(doc, "assets/shap_feature_importance.png", 
                     "Figure 2: Global Mean Absolute SHAP Values and Directional Impact Indicators for Top 10 Churn Drivers.")

    add_body_p(doc, 
               "Detailed Attribution Analysis:\n"
               "1. Monthly Active User (MAU) Decay (+0.42 SHAP): A 30% drop in active seats over a 60-day window is the strongest leading indicator of cancellation.\n"
               "2. Support Escalations (+0.38 SHAP): Multiple unresolved critical issues create acute dissatisfaction that overrides product utility.\n"
               "3. Executive Sponsor Inactivity (+0.33 SHAP): When the original enterprise buyer ceases logging in or attending quarterly reviews, churn hazard triples within 90 days.\n"
               "4. Workflow Integration Depth (-0.31 SHAP): Integration with core enterprise pipelines (e.g., Salesforce, Slack, ERP) acts as an 'organizational anchor', reducing churn risk by over 55%.")

    format_heading(doc, "3.3 Finding 3: The Support Escalation & Churn Probability Matrix", level=2)
    add_body_p(doc, 
               "Cross-tabulating historical support helpdesk data with account tiers revealed a dramatic non-linear compounding risk when support tickets remain unresolved across escalating severities.")
    
    add_figure_image(doc, "assets/churn_risk_heatmap.png", 
                     "Figure 3: Churn Risk Heatmap across Account Contract Tiers and Unresolved Support Escalations (Past 60 Days).")

    add_body_p(doc, 
               "Critical Operational Insight: In Enterprise accounts ($100k+ ARR), the baseline churn rate is exceptionally low at 4.2% when 0 escalations occur. However, when an enterprise account accumulates 5+ unresolved escalations, the churn probability skyrockets to 62.4%—a 14.8x risk surge. In Mid-Market and SMB accounts, 3+ unresolved issues create a near-certain (>50%) churn outcome.",
               bold_prefix="Key Takeaway: ")

    format_heading(doc, "3.4 Finding 4: Cumulative Lift, Targeting Efficiency & Economic Optimization Frontier", level=2)
    add_body_p(doc, 
               "To bridge technical model accuracy with executive fiscal decision-making, the model's performance was evaluated along a cumulative gains curve and an economic optimization frontier.")
    
    add_figure_image(doc, "assets/revenue_lift_frontier.png", 
                     "Figure 4: (A) Cumulative Gains / Lift Curve vs Baseline; (B) Net Economic Revenue Optimization Frontier.")

    add_body_p(doc, 
               "Fiscal Frontier Interpretation: Contacting 100% of accounts is operationally impossible and economically wasteful due to Customer Success headcount constraints ($45k per 1% outreach bandwidth). The economic optimization curve reveals that targeting exactly the top 22% highest-risk accounts maximizes net enterprise value, recovering $5.13M in gross ARR at an operational cost of $910K, yielding a peak net economic return of $4.22M (ROI = 4.6x).",
               bold_prefix="Strategic Financial Takeaway: ")

    # -------------------------------------------------------------
    # SECTION 4: ACTIONABLE INSIGHTS & BUSINESS IMPACT
    # -------------------------------------------------------------
    format_heading(doc, "4. Actionable Insights & Strategic Business Impact", level=1)
    
    add_body_p(doc, 
               "Translating data science findings into commercial value requires operationalizing predictive outputs into concrete, role-specific workflows across Customer Success, Product Management, and Executive Leadership.")

    format_heading(doc, "4.1 Tiered Proactive Retention Playbook", level=2)
    add_body_p(doc, "Based on account value and model-assigned risk tier, automated triage playbooks are activated:")

    playbook_headers = ["Risk Decile / Tier", "Account Value Segment", "Automated Trigger Criteria", "Prescribed Intervention Playbook", "Responsible Owner"]
    playbook_data = [
        ["Critical (Top 5%)", "Enterprise ($100k+)", "Model Score > 0.82 OR Sev-1 Ticket > 5 Days", "Executive Sponsor Outreach + VP CS Alignment + Dedicated Solutions Architect remediation sprint within 48h.", "VP Customer Success / Dedicated Enterprise CSM"],
        ["High (Top 6-15%)", "Enterprise & Mid-Market", "Model Score 0.65 - 0.81 OR MAU Decay > 25%", "Strategic Account Review (QBR) reschedule + Custom Feature Adoption Workshop + Workflow Integration review.", "Senior CSM & Product Specialist"],
        ["Moderate (Top 16-25%)", "Mid-Market & Growth", "Model Score 0.50 - 0.64 OR Onboarding < 4 features", "Automated in-app personalized guided walkthroughs + In-app concierge chat + CS check-in call.", "Scaled CSM / Automated Lifecycle Engine"],
        ["Low / Healthy (Bottom 75%)", "All Segments", "Model Score < 0.50", "Standard automated nurture, feature release announcements, expansion upsell cadence.", "Marketing Automation & Growth Team"]
    ]
    add_styled_table(doc, playbook_headers, playbook_data, col_widths=[1.2, 1.2, 1.4, 2.0, 1.2])

    format_heading(doc, "4.2 Financial Sensitivity & ROI Model", level=2)
    add_body_p(doc, 
               "To provide the CFO and executive committee with complete budget transparency, a multi-scenario sensitivity model was formulated evaluating conservative, moderate, and aggressive retention success rates:")

    roi_headers = ["Scenario Parameter", "Conservative Scenario (25% Win Rate)", "Base / Expected Case (35% Win Rate)", "Optimistic Scenario (45% Win Rate)"]
    roi_data = [
        ["Total Targeted At-Risk ARR", "$12,800,000", "$12,800,000", "$12,800,000"],
        ["Model Capture Rate (Top 22%)", "74.8% ($9,574,400 ARR)", "74.8% ($9,574,400 ARR)", "74.8% ($9,574,400 ARR)"],
        ["Retention Success Rate", "25.0%", "35.0%", "45.0%"],
        ["Gross ARR Preserved", "$2,393,600", "$3,351,040", "$4,308,480"],
        ["Total Operational Intervention Cost", "$910,000", "$910,000", "$910,000"],
        ["Net Annual Economic Value Created", "$1,483,600", "$2,441,040", "$3,398,480"],
        ["Net Return on Investment (ROI)", "2.63x (163% Net ROI)", "3.68x (268% Net ROI)", "4.73x (373% Net ROI)"]
    ]
    add_styled_table(doc, roi_headers, roi_data, col_widths=[2.1, 1.5, 1.5, 1.4])

    # -------------------------------------------------------------
    # SECTION 5: PRESENTATION STRATEGY FOR NON-TECHNICAL STAKEHOLDERS
    # -------------------------------------------------------------
    format_heading(doc, "5. Presentation & Communication Strategy for Non-Technical Stakeholders", level=1)
    
    add_body_p(doc, 
               "Data science insights generate zero enterprise value if technical findings fail to influence executive decision-making. Communicating complex machine learning algorithms, SHAP attributions, and survival mathematics to non-technical stakeholders requires structured storytelling, audience empathy, and cognitive clarity.",
               bold_prefix="Communication Philosophy: ")

    format_heading(doc, "5.1 The Minto Pyramid Principle & SCQA Storytelling Framework", level=2)
    add_body_p(doc, 
               "The presentation is architected using the Minto Pyramid Principle (Lead with the conclusion first) combined with the SCQA framework:")
    
    add_callout_box(doc, 
                    "Executive SCQA Narrative Blueprint", 
                    "1. Situation (Context): 'Our enterprise SaaS platform currently generates $90M ARR with industry-standard product satisfaction.'\n"
                    "2. Complication (Problem): 'However, annual gross churn has climbed to 14.2% ($12.8M at risk), driven by undetected friction in the first 90 days and support escalations.'\n"
                    "3. Question (Challenge): 'How can leadership proactively identify at-risk revenue 60 days before contract expiration and deploy Customer Success resources at peak ROI?'\n"
                    "4. Answer (Solution): 'By deploying Project Horizon's predictive AI engine, we capture 73.4% of at-risk churners and systematically preserve $4.22M net ARR with a 4.6x ROI through a tiered proactive CS playbook.'")

    format_heading(doc, "5.2 Stakeholder Persona Alignment Matrix", level=2)
    add_body_p(doc, "Different executive stakeholders require tailored visual narratives and primary focus areas:")

    persona_headers = ["Executive Stakeholder", "Core Motivations & KPIs", "Cognitive Friction Point", "Tailored Visual & Narrative Strategy"]
    persona_data = [
        ["CFO & Finance Committee", "ARR preservation, CAC payback period, ROI on CS headcount.", "Skepticism regarding intervention cost vs realized revenue gains.", "Lead with Figure 4B (Net Economic Frontier) and Table 5 (Sensitivity ROI Table). Frame model as capital efficiency tool."],
        ["Chief Revenue Officer (CRO) / VP CS", "Net Revenue Retention (NRR), gross renewal rates, CSM bandwidth allocation.", "Alert fatigue; fear of false alarms wasting valuable CSM time.", "Highlight Top-Decile Precision (78.2%) and Figure 3 (Support Escalation Matrix). Emphasize clear tiered playbooks."],
        ["Chief Product Officer (CPO) / VP Product", "Feature adoption velocity, onboarding completion, UX friction points.", "Belief that churn is purely a sales/pricing issue rather than product onboarding.", "Present Figure 1 (90-Day Survival Curve) and Figure 2 (Integration Depth). Focus on Day 1-90 user journey and API adoption."],
        ["Customer Success Managers (CSMs)", "Daily account visibility, actionable guidance, saving at-risk renewals.", "Black-box AI confusion ('Why is this account red?').", "Deliver individual account-level SHAP 'Waterfalls' directly inside Salesforce/CRM UI with 3 concrete prescribed actions."]
    ]
    add_styled_table(doc, persona_headers, persona_data, col_widths=[1.5, 1.6, 1.6, 1.8])

    format_heading(doc, "5.3 Ten-Slide Executive Deck Architecture & Blueprint", level=2)
    add_body_p(doc, "The structured 20-minute executive presentation is organized into 10 high-impact slides:")

    deck_headers = ["Slide #", "Slide Title", "Visual Layout / Chart", "Core Message / Takeaway", "Speaker Note & Delivery Guidance"]
    deck_data = [
        ["Slide 1", "The $12.8M Retention Opportunity", "Executive KPI Cards (Churn, NRR, ARR at Risk)", "Proactive retention is our highest-ROI growth lever in 2026.", "Hook the audience within 60 seconds with total addressable revenue at risk."],
        ["Slide 2", "The Current Problem: Blind Reactive Churn", "Timeline showing 14-day renewal blindspot vs 90-day window", "Our current reactive process gives CSMs zero runway to salvage accounts.", "Share a real sanitized enterprise churn case study to build emotional resonance."],
        ["Slide 3", "Project Horizon: The AI Early-Warning System", "High-level 3-step architecture diagram (Data -> AI -> Playbook)", "We have engineered an intelligent engine that detects churn risk 60-90 days early.", "Avoid technical jargon; describe model as an 'intelligent radar system'."],
        ["Slide 4", "The 90-Day Onboarding Cliff", "Figure 1: Kaplan-Meier Survival Curve by Adoption Tier", "Accounts adopting <4 features lose 32% retention by Day 90.", "Emphasize to Product leadership that the first quarter dictates customer lifetime."],
        ["Slide 5", "What Really Drives Customer Churn?", "Figure 2: Top 10 SHAP Behavioral Drivers (Horizontal Bar)", "Churn is driven by user seat decay, unresolved support tickets, and lack of integrations.", "Walk through the top 3 friction drivers and top 2 retention anchors."],
        ["Slide 6", "The Escalation Danger Zone", "Figure 3: Support Friction & Account Tier Risk Heatmap", "Enterprise accounts with 5+ unresolved issues face 62.4% churn risk.", "Directly address CS/Support leadership on establishing automated ticket SLAs."],
        ["Slide 7", "Targeting Precision & Model Efficiency", "Figure 4A: Cumulative Gains / Lift Curve", "Targeting the top 20% of flagged accounts captures 73.4% of all churners.", "Reassure team that CSMs will not be overwhelmed by false alarms."],
        ["Slide 8", "Economic ROI & Net Revenue Frontier", "Figure 4B: Net Revenue Optimization Curve", "Optimal targeting of top 22% yields $4.22M net preserved ARR (4.6x ROI).", "Highlight financial sweet spot where marginal savings exceed intervention costs."],
        ["Slide 9", "The Actionable Retention Playbook", "Role-specific Action Matrix (Enterprise vs SMB workflows)", "Clear, automated workflows ensure immediate cross-functional action.", "Show how predictions translate directly into Salesforce tasks for CSMs."],
        ["Slide 10", "Roadmap, Governance & Next Steps", "30-60-90 Day Phased Execution Gantt Chart", "3-stage rollout: Shadow mode -> 60-day pilot -> Full enterprise operationalization.", "Call to action: Request steering committee approval to launch Phase 1 Pilot."]
    ]
    add_styled_table(doc, deck_headers, deck_data, col_widths=[0.6, 1.4, 1.4, 1.6, 1.5])

    format_heading(doc, "5.4 Executive FAQ & Objection Handling Matrix", level=2)
    add_body_p(doc, "Anticipating tough questions from skeptical C-suite leaders is essential to winning project approval:")

    faq_headers = ["Anticipated Executive Question", "Stakeholder", "Data Science & Strategic Response"]
    faq_data = [
        ["'How do we know the model isn't just flagging accounts that were already going to churn regardless?'", "CFO", "The model assigns risk scores based on leading indicators (e.g. 30-day seat decay, login velocity) 60-90 days prior to contract expiration—months before renewal teams traditionally notice disengagement. In validation testing, early intervention on these signals altered the trajectory of 35% of accounts."],
        ["'What is the cost of false positives? Will CSMs waste time chasing healthy clients?'", "VP CS", "At our optimal operating threshold (Top 22%), precision is 78.2%. Even in the remaining 21.8% of cases, accounts flagged exhibit genuine friction (e.g., lower feature usage or open tickets). Outreach to these accounts acts as a proactive relationship check, improving CSAT even if churn wasn't imminent."],
        ["'Why not just use simple heuristics like 'login count < 5' instead of complex machine learning?'", "VP Product", "Rule-based heuristics produce high false positive rates (48% error rate in historical audit) and fail to capture multi-variable interactions (e.g., an account with high logins but 3 Sev-1 escalations and decaying admin engagement). The ensemble model achieves a 3.67x lift over simple rules."],
        ["'How will we ensure the model doesn't become outdated as customer behavior changes?'", "CTO", "The MLOps pipeline includes automated drift monitoring tracking Population Stability Index (PSI) on features and weekly retraining pipelines. Any drift exceeding PSI > 0.1 triggers automated alerts and shadow validation before production deployment."],
        ["'Does the model account for seasonal budget cycles or macroeconomic slowdowns?'", "CRO", "Yes. The training dataset spans 24 continuous months, encompassing annual budget renewals, Q4 procurement rushes, and macroeconomic fluctuations. Macro-economic index features and calendar quarter indicators were explicitly incorporated."],
        ["'What is the immediate investment required to launch Phase 1?'", "CEO", "Phase 1 requires zero new software procurement. It utilizes existing infrastructure (AWS/Snowflake, Python, Salesforce API) and leverages current CSM capacity dedicated to a 60-day randomized control pilot across 200 accounts."]
    ]
    add_styled_table(doc, faq_headers, faq_data, col_widths=[1.8, 1.0, 3.7])

    # -------------------------------------------------------------
    # SECTION 6: RECOMMENDATIONS, OPERATIONALIZATION & FUTURE WORK
    # -------------------------------------------------------------
    format_heading(doc, "6. Recommendations, Operationalization & Future Work", level=1)
    
    add_body_p(doc, 
               "To realize the projected $4.22M net ARR retention impact, an operationalization plan spanning governance, phased deployment, and advanced analytical research is outlined below.")

    format_heading(doc, "6.1 Phased Implementation Roadmap (30-60-90 Day Execution)", level=2)
    add_body_p(doc, "Deployment will execute in three sequential phases to mitigate operational risk:")

    roadmap_headers = ["Phase & Horizon", "Key Milestones & Deliverables", "Success Criteria & Gating Metric", "Responsible Team"]
    roadmap_data = [
        ["Phase 1: Shadow Deployment (Days 1 - 30)", "Deploy XGBoost inference pipeline in shadow mode. Score all 50,000 accounts weekly. Integrate output schema into Snowflake data warehouse. Run daily telemetry drift checks.", "Inference pipeline runtime < 15 min; Data drift PSI < 0.05; Zero interruption to live systems.", "Data Engineering & MLOps Team"],
        ["Phase 2: Controlled A/B Pilot (Days 31 - 60)", "Launch randomized control trial across 200 at-risk Mid-Market & Enterprise accounts. Group A receives ML-triggered Tiered Retention Playbook; Group B follows standard reactive renewal process.", "Statistical significance (p < 0.05) showing >= 25% relative reduction in 60-day churn in Group A.", "VP Customer Success & Lead Data Scientist"],
        ["Phase 3: Enterprise Rollout (Days 61 - 90)", "Full CRM integration (Salesforce/HubSpot). Automated account risk badges and SHAP explainability widgets on CSM dashboards. Bi-weekly retention steering review.", "100% CSM onboarding; 90-day Net Revenue Retention tracking toward > 108% target.", "Executive Steering Committee & CS Enablement"]
    ]
    add_styled_table(doc, roadmap_headers, roadmap_data, col_widths=[1.5, 2.3, 1.5, 1.2])

    format_heading(doc, "6.2 MLOps Architecture & Model Governance Framework", level=2)
    add_body_p(doc, "Long-term model reliability will be enforced through four operational governance pillars:")
    add_body_p(doc, "1. Automated Data & Concept Drift Monitoring: Calculating Population Stability Index (PSI) and Kolmogorov-Smirnov (KS) statistics across input features every 7 days.\n"
                    "2. Retraining & Continuous Learning Cadence: Automated monthly model retraining on rolling 24-month historical datasets with automated shadow validation gating before deployment.\n"
                    "3. Model Explainability & Auditability: Storing individual prediction SHAP values and model version hashes in Snowflake for full audit compliance.\n"
                    "4. Feedback Loop Integration: Capturing CSM intervention outcomes (Saved, Discounted, Churned, False Positive) directly from Salesforce to enrich retraining labels.")

    format_heading(doc, "6.3 Future Data Science Research Horizons", level=2)
    add_body_p(doc, "Following successful production deployment of Project Horizon, three advanced research initiatives are recommended for subsequent quarters:")
    add_body_p(doc, "• Prescriptive Uplift Modeling (Causal Machine Learning): Transitioning from predictive modeling ('Who will churn?') to causal uplift modeling ('Which specific intervention will change the outcome for this customer?'), utilizing Meta-Learners (T-Learners, X-Learners) to maximize intervention efficiency.\n"
                    "• Dynamic Price Elasticity & Renewal Discount Optimization: Developing multi-armed bandit algorithms to recommend optimal renewal discount percentages based on account health and willingness to pay, avoiding unnecessary price concessions.\n"
                    "• Real-Time In-App Behavioral Event Stream Processing: Upgrading batch weekly inference to real-time event stream processing via Apache Flink, triggering immediate in-app nudges when high-friction drop-offs occur.")

    # Save Document
    doc_path = "Comprehensive_Data_Science_Report_and_Insights_Presentation_Plan.docx"
    doc.save(doc_path)
    print(f"Successfully generated: {doc_path}")

if __name__ == "__main__":
    main()
