"""
build_ml_plan_doc.py
Generates an enterprise-grade, highly structured, comprehensive Word Document (.docx)
for Week 3 Task: Python-Based Machine Learning Model Development and Evaluation Plan.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

base_dir = os.path.dirname(os.path.abspath(__file__))
doc_path = os.path.join(base_dir, "ML_Model_Development_and_Evaluation_Plan.docx")

# Color Palette Constants
COLOR_NAVY = RGBColor(27, 54, 93)       # #1B365D - Primary
COLOR_STEEL = RGBColor(46, 107, 158)    # #2E6B9E - Secondary
COLOR_CHARCOAL = RGBColor(44, 62, 80)   # #2C3E50 - Body
COLOR_TEAL = RGBColor(0, 128, 128)      # #008080 - Accent
COLOR_MUTED = RGBColor(100, 110, 120)   # #646E78 - Subtitles / Notes
HEX_NAVY = "1B365D"
HEX_STEEL = "2E6B9E"
HEX_LIGHT_BG = "F4F7FA"
HEX_CALLOUT_BG = "EEF4F8"
HEX_ALERT_BG = "FFF8E7"
HEX_BORDER = "CCCCCC"

def set_cell_background(cell, hex_color):
    """Sets background shading color for a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_header_footer(doc):
    """Adds running headers and footers."""
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        
        # Header
        hdr = s.header
        hdr_p = hdr.paragraphs[0]
        hdr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hdr_p.add_run("Virtual Data Science Explorer Intern | Week 3: ML Model Development & Evaluation Plan")
        hrun.font.name = 'Calibri'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = COLOR_MUTED
        
        # Footer
        ftr = s.footer
        ftr_p = ftr.paragraphs[0]
        ftr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = ftr_p.add_run("Confidential & Proprietary — Python-Based Machine Learning Architecture Blueprint")
        frun.font.name = 'Calibri'
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = COLOR_MUTED

def style_paragraph(p, font_name='Calibri', size_pt=11, color=COLOR_CHARCOAL, line_spacing=1.15, space_after=6, bold=False, italic=False):
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    for run in p.runs:
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.color.rgb = color
        run.font.bold = bold
        run.font.italic = italic

def add_heading_1(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = COLOR_NAVY
    return h

def add_heading_2(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COLOR_STEEL
    return h

def add_heading_3(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(2)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_CHARCOAL
    return h

def add_body_p(doc, text, bold_prefix=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(space_after)
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_NAVY
    r_body = p.add_run(text)
    r_body.font.name = 'Calibri'
    r_body.font.size = Pt(10.5)
    r_body.font.color.rgb = COLOR_CHARCOAL
    return p

def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * level + 0.25)
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_NAVY
    r_body = p.add_run(text)
    r_body.font.name = 'Calibri'
    r_body.font.size = Pt(10.5)
    r_body.font.color.rgb = COLOR_CHARCOAL
    return p

def add_callout(doc, title, body_text, hex_bg=HEX_CALLOUT_BG, border_color="1B365D"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, hex_bg)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    # Border
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    r_title = p.add_run(title + "\n")
    r_title.font.name = 'Segoe UI'
    r_title.font.size = Pt(10.5)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_NAVY
    
    r_body = p.add_run(body_text)
    r_body.font.name = 'Calibri'
    r_body.font.size = Pt(10)
    r_body.font.color.rgb = COLOR_CHARCOAL
    
    # spacing after table
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

def add_figure(doc, img_name, caption_text, width_inches=6.3):
    img_path = os.path.join(base_dir, img_name)
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(4)
        run = p_img.add_run()
        run.add_picture(img_path, width=Inches(width_inches))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        r_cap = p_cap.add_run(f"Figure: {caption_text}")
        r_cap.font.name = 'Calibri'
        r_cap.font.size = Pt(9.5)
        r_cap.font.italic = True
        r_cap.font.color.rgb = COLOR_MUTED

def style_table(doc, tbl, col_widths, headers, data):
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr_row = tbl.rows[0]
    for idx, heading in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.width = Inches(col_widths[idx])
        set_cell_background(cell, HEX_NAVY)
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(heading)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    # Data rows
    for r_idx, row_data in enumerate(data):
        row = tbl.add_row()
        bg_hex = HEX_LIGHT_BG if (r_idx % 2 == 1) else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = Inches(col_widths[c_idx])
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            run.font.color.rgb = COLOR_CHARCOAL

    # Set table borders
    tblPr = tbl._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                <w:insideV w:val="none"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
            </w:tblBorders>
        ''')
        tblPr[0].append(borders)
        
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(0)
    p_sp.paragraph_format.space_after = Pt(6)

def build_document():
    doc = Document()
    add_header_footer(doc)
    
    # -------------------------------------------------------------
    # Cover Page / Header Block
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(4)
    r_t = title_p.add_run("PYTHON-BASED MACHINE LEARNING MODEL DEVELOPMENT & EVALUATION PLAN")
    r_t.font.name = 'Segoe UI'
    r_t.font.size = Pt(22)
    r_t.font.bold = True
    r_t.font.color.rgb = COLOR_NAVY
    
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(16)
    r_sub = sub_p.add_run("An Enterprise Architecture Blueprint for Predictive Churn Modeling, Rigorous Validation, and Production MLOps Serving")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = COLOR_STEEL

    # Metadata Table Box
    meta_tbl = doc.add_table(rows=4, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Program & Role:", "Virtual Data Science Explorer Intern — Machine Learning Track"),
        ("Milestone / Task:", "Week 3 Task: ML Model Development and Evaluation Plan"),
        ("Author / Intern:", "Data Science & MLOps Engineering Intern"),
        ("Estimated Dev Effort:", "30 – 35 Hours Comprehensive Lifecycle Blueprint")
    ]
    for idx, (k, v) in enumerate(meta_data):
        c0 = meta_tbl.rows[idx].cells[0]
        c1 = meta_tbl.rows[idx].cells[1]
        c0.width = Inches(2.2)
        c1.width = Inches(4.3)
        set_cell_background(c0, HEX_LIGHT_BG)
        set_cell_background(c1, HEX_LIGHT_BG)
        set_cell_margins(c0, 60, 60, 100, 100)
        set_cell_margins(c1, 60, 60, 100, 100)
        
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(k)
        r0.font.name = 'Segoe UI'
        r0.font.size = Pt(9.5)
        r0.font.bold = True
        r0.font.color.rgb = COLOR_NAVY
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(v)
        r1.font.name = 'Calibri'
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = COLOR_CHARCOAL

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Executive Callout
    add_callout(doc, "Executive Summary & Architectural Scope",
                "This document establishes an end-to-end, industry-grade plan for engineering, benchmarking, evaluating, and deploying a high-performance machine learning model in Python. Addressing a multi-million-dollar subscription customer churn problem, this specification provides mathematically rigorous preprocessing pipelines, multi-algorithm selection protocols, Bayesian hyperparameter optimization, cost-benefit validation frameworks, model explainability (SHAP), and a resilient MLOps deployment architecture.")

    # -------------------------------------------------------------
    # SECTION 1: PROBLEM DEFINITION & BUSINESS FORMULATION
    # -------------------------------------------------------------
    add_heading_1(doc, "1. Problem Definition & Business Formulation")
    
    add_heading_2(doc, "1.1 Hypothetical Business Scenario: Enterprise SaaS Churn Prediction")
    add_body_p(doc, "In subscription-based SaaS and telecommunications enterprises, customer acquisition costs (CAC) typically exceed customer retention costs by 5x to 7x. Retaining high-value accounts directly compounds monthly recurring revenue (MRR) and customer lifetime value (LTV).")
    add_body_p(doc, "The objective of this project is to build an automated, real-time Machine Learning scoring engine in Python that predicts customer churn risk (probability of cancellation within a 60-day forward-looking window) based on telemetry, billing patterns, support interaction frequency, and contract characteristics.")

    add_heading_2(doc, "1.2 Technical Formulation & Mathematical Objectives")
    add_body_p(doc, "The challenge is formulated as a supervised binary classification problem with probability calibration:")
    add_bullet(doc, "Feature Vector X ∈ ℝᴰ: Vector containing D preprocessed demographic, behavioral, contractual, and operational features.", bold_prefix="Input Space: ")
    add_bullet(doc, "Target Label y ∈ {0, 1}: Where y = 1 denotes a churn event within 60 days, and y = 0 denotes an active/retained subscriber.", bold_prefix="Output Space: ")
    add_bullet(doc, "Learned Hypothesis P(y = 1 | X): A calibrated posterior probability indicating the likelihood of churn.", bold_prefix="Objective Function: ")

    add_heading_2(doc, "1.3 Business KPIs vs. Machine Learning Technical Metrics")
    add_body_p(doc, "To bridge the gap between engineering metrics and financial impact, the development plan maps ML technical performance directly to enterprise business outcomes:")
    
    kpi_headers = ["Domain Level", "Primary Metric", "Target Threshold", "Business / Financial Impact"]
    kpi_data = [
        ["Business KPI", "Customer Churn Rate", "Reduction from 18.5% to < 12.0%", "Saves ~$3.4M annually in preserved recurring subscription revenue."],
        ["Business KPI", "Retention Campaign ROI", "≥ 4.5x Return on Retention Spend", "Ensures high-cost incentives (discounts/concierge) target true at-risk accounts."],
        ["Technical ML", "Precision-Recall AUC (PR-AUC)", "PR-AUC ≥ 0.68 (Baseline: 0.18)", "Provides superior discrimination on imbalanced tabular distribution."],
        ["Technical ML", "Recall @ Top 20% Decile", "Recall@20 ≥ 75.0%", "Guarantees top retention outreach captures 3 out of every 4 churners."],
        ["Technical ML", "Brier Score (Calibration)", "Brier Score ≤ 0.09", "Ensures output probabilities reflect true empirical frequencies for tiered risk triage."]
    ]
    tbl_kpi = doc.add_table(rows=1, cols=4)
    style_table(doc, tbl_kpi, [1.3, 1.6, 1.6, 2.0], kpi_headers, kpi_data)

    # -------------------------------------------------------------
    # SECTION 2: END-TO-END ML WORKFLOW ARCHITECTURE
    # -------------------------------------------------------------
    add_heading_1(doc, "2. End-to-End Machine Learning Workflow Architecture")
    add_body_p(doc, "Building an enterprise-ready ML solution requires a modular, reproducible, and leak-free pipeline across 7 distinct lifecycle phases. Figure 1 illustrates the full architecture from ingestion to automated continuous retraining.")

    add_figure(doc, "workflow_architecture.png", "End-to-End Machine Learning System Lifecycle Architecture and Feedback Loops.")

    add_body_p(doc, "The lifecycle is structured into distinct, decoupled stages to ensure continuous integration, testing, and governance:")
    add_bullet(doc, "Automated data ingestion from telemetry lakes, billing databases, and CRM logs with Pydantic/Great Expectations schema validation.", bold_prefix="Phase 1 (Ingestion & Profiling): ")
    add_bullet(doc, "Strict train-test isolated transformations, outlier treatment, Yeo-Johnson transforms, target encoding, and SMOTE-NC balancing.", bold_prefix="Phase 2 (Preprocessing Pipeline): ")
    add_bullet(doc, "Benchmarking Logistic Regression, Random Forest, XGBoost, LightGBM, CatBoost, and TabNet architectures.", bold_prefix="Phase 3 (Model Exploration): ")
    add_bullet(doc, "5-Fold Nested Stratified Cross-Validation combined with Optuna Bayesian search under PR-AUC and cost-benefit utility matrices.", bold_prefix="Phase 4 (Validation & Optimization): ")
    add_bullet(doc, "Global & local feature attributions via TreeSHAP, demographic fairness parity tests, and artifact signing.", bold_prefix="Phase 5 (Explainability & Governance): ")
    add_bullet(doc, "Containerized FastAPI microservice with sub-25ms response latency, ONNX runtime optimization, and automated health checks.", bold_prefix="Phase 6 (Production Deployment): ")
    add_bullet(doc, "Real-time Population Stability Index (PSI) tracking, Kolmogorov-Smirnov drift detection, and automated trigger pipelines.", bold_prefix="Phase 7 (MLOps Observability): ")

    # -------------------------------------------------------------
    # SECTION 3: DATA PREPROCESSING & FEATURE ENGINEERING PIPELINE
    # -------------------------------------------------------------
    add_heading_1(doc, "3. Data Preprocessing & Feature Engineering Pipeline")
    add_body_p(doc, "Data quality and feature representation form the foundation of model performance. A flawed preprocessing step introduces subtle data leakage, leading to artificially inflated cross-validation scores that collapse in production.")

    add_figure(doc, "preprocessing_pipeline.png", "Modular Data Preprocessing & Feature Engineering Pipeline with Strict Train-Test Isolation.")

    add_heading_2(doc, "3.1 Data Cleaning & Anomaly Treatment")
    add_bullet(doc, "Diagnosing missingness mechanisms (MCAR, MAR, MNAR). Numerical fields with MAR are imputed via IterativeImputer (MICE) or KNNImputer (k=5) fitted strictly on training folds. Categorical columns receive an explicit 'Missing_Category' token.", bold_prefix="Missing Value Imputation: ")
    add_bullet(doc, "Outliers in usage telemetry (e.g., extreme bandwidth or API call spikes) are flagged using Isolation Forest (contamination=0.01) and winsorized at 1st and 99th percentiles rather than arbitrarily dropped.", bold_prefix="Outlier Detection & Winsorization: ")
    add_bullet(doc, "Enforcing strict schema typing, parsing ISO 8601 timestamps into elapsed tenures, and validating primary keys across CRM and Billing joins.", bold_prefix="Schema & Type Integrity: ")

    add_heading_2(doc, "3.2 Feature Transformation & Scaling")
    add_bullet(doc, "Highly skewed financial metrics (e.g., MonthlySpend, TotalRevenue) undergo Yeo-Johnson power transformations to normalize residual distributions, followed by RobustScaler to resist remaining heavy tails.", bold_prefix="Numerical Transformation: ")
    add_bullet(doc, "Low-cardinality nominal variables (e.g., ContractType, PaymentMethod) are One-Hot Encoded (drop='first'). High-cardinality features (e.g., PostalRegion, IndustrySubcategory) undergo Smoothed Target Encoding with cross-validation regularization to prevent target leakage.", bold_prefix="Categorical Encoding: ")
    add_bullet(doc, "Contract length and customer tenure are segmented into ordinal risk brackets based on empirical hazard rates.", bold_prefix="Ordinal Binning: ")

    add_heading_2(doc, "3.3 Advanced Domain Feature Engineering")
    add_body_p(doc, "Raw tabular features are enriched through mathematical transformations that capture behavioral decay and usage velocity:")
    add_bullet(doc, "Ratio of usage volume in the last 30 days compared to the 90-day moving average. A ratio < 0.70 signals rapid disengagement.", bold_prefix="1. Usage Decay Velocity (ΔUsage_30_90): ")
    add_bullet(doc, "Frequency of high-severity CRM tickets logged per month of active tenure. Captures unresolved friction points.", bold_prefix="2. Support Friction Index (Tickets / Tenure): ")
    add_bullet(doc, "Rolling standard deviation of invoice amounts divided by the mean, identifying billing disputes and erratic charges.", bold_prefix="3. Payment Volatility Index (CV_Billing): ")
    add_bullet(doc, "Multiplicative interaction between tenure duration and discount percentage, distinguishing loyalty discounts from short-term promo seekers.", bold_prefix="4. Loyalty-Discount Interaction (Tenure × Discount%): ")

    add_heading_2(doc, "3.4 Feature Selection & Dimensionality Management")
    add_body_p(doc, "To eliminate noise and avoid the curse of dimensionality, a multi-stage filtering protocol is applied:")
    add_bullet(doc, "Features with zero or near-zero variance (variance < 0.01) are immediately removed.", bold_prefix="Stage 1 (Variance Thresholding): ")
    add_bullet(doc, "Pairwise Pearson/Spearman correlation matrices and Variance Inflation Factor (VIF > 5.0) eliminate collinear redundant features.", bold_prefix="Stage 2 (Collinearity Pruning): ")
    add_bullet(doc, "Calculating Mutual Information (MI) against target y to identify non-linear relationships.", bold_prefix="Stage 3 (Non-Linear Mutual Information): ")
    add_bullet(doc, "Recursive Feature Elimination with Stratified 5-Fold Cross-Validation (RFECV) on a LightGBM estimator to select the optimal subset of predictive features.", bold_prefix="Stage 4 (RFECV Wrapper Selection): ")

    add_heading_2(doc, "3.5 Addressing Severe Class Imbalance")
    add_body_p(doc, "Customer churn typically exhibits a 15% to 20% minority class proportion. Standard accuracy-driven training collapses towards the majority class. The plan employs a 3-pronged mitigation strategy:")
    add_bullet(doc, "Synthetic Minority Over-sampling Technique for Nominal and Continuous features applied EXCLUSIVELY to training folds (never to validation or test splits).", bold_prefix="1. Resampling via SMOTE-NC: ")
    add_bullet(doc, "Incorporating scale_pos_weight = (N_neg / N_pos) in LightGBM/XGBoost, penalizing false negative errors proportionally during gradient descent.", bold_prefix="2. Cost-Weighted Loss Optimization: ")
    add_bullet(doc, "Adjusting the classification decision threshold from default 0.50 down to optimal probability threshold θ* derived via cost-benefit utility curve optimization.", bold_prefix="3. Post-Hoc Threshold Optimization: ")

    # -------------------------------------------------------------
    # SECTION 4: MODEL SELECTION, ARCHITECTURE & TRAINING
    # -------------------------------------------------------------
    add_heading_1(doc, "4. Model Selection, Architecture Exploration & Training")
    add_body_p(doc, "To ensure optimal generalization on structured tabular data, a systematic model benchmark is planned across distinct model families, ranging from linear baselines to state-of-the-art gradient boosted trees and deep tabular architectures.")

    add_heading_2(doc, "4.1 Candidate Model Comparison Matrix")
    
    model_headers = ["Model Family", "Algorithm", "Key Strengths", "Limitations & Risks", "Suitability Rank"]
    model_data = [
        ["Linear Baseline", "Logistic Regression (L1/L2 ElasticNet)", "Highly interpretable, fast inference (<1ms), clear benchmark.", "Fails to capture complex non-linear feature interactions without manual polynomial expansion.", "Baseline / Reference (Rank 4)"],
        ["Ensemble Bagging", "Random Forest Classifier", "Robust to outliers, captures non-linearities, low risk of overfitting.", "High memory footprint, slower prediction throughput on large trees, lacks native categorical handling.", "Candidate (Rank 3)"],
        ["Gradient Boosting", "LightGBM / XGBoost / CatBoost", "State-of-the-art tabular performance, native categorical support (LightGBM/CatBoost), fast GPU training, robust regularization.", "Requires careful hyperparameter tuning; can overfit on noisy small datasets if unregularized.", "Primary Champion Candidate (Rank 1)"],
        ["Deep Learning", "TabNet / Multi-Layer Perceptron", "Sequential attention mechanism, handles multi-modal inputs, end-to-end embedding learning.", "High computational cost, slower inference latency, requires massive data volume to surpass gradient boosting.", "Challenger (Rank 2)"]
    ]
    tbl_models = doc.add_table(rows=1, cols=5)
    style_table(doc, tbl_models, [1.1, 1.3, 1.6, 1.6, 0.9], model_headers, model_data)

    add_heading_2(doc, "4.2 Justification for Champion Architecture: LightGBM / XGBoost Ensemble")
    add_body_p(doc, "Gradient Boosted Decision Trees (specifically LightGBM and CatBoost) consistently outperform deep neural networks on tabular datasets characterized by heterogeneous feature types, varying scales, and non-smooth decision boundaries. Their Histogram-based split finding (GOSS in LightGBM) provides extreme computational efficiency, native handling of missing values, and built-in L1/L2 leaf regularization.")
    add_body_p(doc, "The proposed plan implements a Stacking Ensemble combining a tuned LightGBM, a tuned CatBoost, and a tuned XGBoost model, with a calibrated Logistic Regression meta-learner for final probability consensus.")

    add_heading_2(doc, "4.3 Hyperparameter Optimization (HPO) Strategy")
    add_body_p(doc, "Traditional Grid Search suffers from exponential combinatorial explosion (curse of dimensionality), while Random Search lacks direction. The development plan utilizes Bayesian Optimization via Optuna:")
    add_bullet(doc, "Models hyperparameter probability distributions p(score | params) and iteratively selects candidate parameters maximizing the Expected Improvement (EI).", bold_prefix="Tree-structured Parzen Estimator (TPE): ")
    add_bullet(doc, "Optuna's MedianPruner and Hyperband prune unpromising trials at early boosting iterations, achieving 5x faster search convergence.", bold_prefix="Asynchronous Trial Pruning: ")
    add_bullet(doc, "Optuna optimizes directly for Out-Of-Fold PR-AUC under 5-Fold Stratified Cross-Validation across 150 optimization trials.", bold_prefix="Search Objective: ")

    hpo_headers = ["Hyperparameter", "Search Range / Distribution", "Optimization Purpose"]
    hpo_data = [
        ["learning_rate (eta)", "Log-Uniform [0.005, 0.20]", "Controls step size shrinkage; smaller rates with early stopping prevent overshooting."],
        ["num_leaves / max_depth", "Integer [15, 127] / [3, 10]", "Governs tree complexity and structural capacity."],
        ["min_child_samples", "Integer [10, 100]", "Prevents leaf nodes from isolating tiny noise clusters (overfitting control)."],
        ["subsample / colsample_bytree", "Uniform [0.50, 0.95]", "Stochastic row/column subsampling to inject bagging diversity."],
        ["reg_alpha (L1) & reg_lambda (L2)", "Log-Uniform [1e-3, 10.0]", "Direct weight penalty enforcing feature sparsity and shrinkage."]
    ]
    tbl_hpo = doc.add_table(rows=1, cols=3)
    style_table(doc, tbl_hpo, [1.8, 2.0, 2.7], hpo_headers, hpo_data)

    # -------------------------------------------------------------
    # SECTION 5: ROBUST VALIDATION & PERFORMANCE EVALUATION
    # -------------------------------------------------------------
    add_heading_1(doc, "5. Robust Validation Strategies & Performance Evaluation")
    add_body_p(doc, "To ensure that model evaluation reflects real-world generalization without data leakage or optimistic bias, the plan establishes a rigorous validation protocol.")

    add_figure(doc, "validation_tuning_flow.png", "Nested Cross-Validation and Bayesian Optimization Workflow with Out-of-Time Test Holdout.")

    add_heading_2(doc, "5.1 Validation Schemes & Leakage Prevention")
    add_bullet(doc, "A full 20% holdout partition sampled strictly from the most recent calendar period (e.g., Q4 data) is isolated at project inception. This dataset is never exposed during preprocessing, feature selection, or hyperparameter optimization.", bold_prefix="1. Out-of-Time (Temporal) Holdout Partition: ")
    add_bullet(doc, "The remaining 80% development dataset is evaluated using 5-Fold Stratified Nested CV. The Outer Loop (5 folds) estimates unbiased generalization error, while the Inner Loop (3 folds) performs Optuna hyperparameter tuning.", bold_prefix="2. 5-Fold Stratified Nested Cross-Validation: ")
    add_bullet(doc, "All transformations (scalers, encoders, imputers, SMOTE) are fit strictly on the training folds and only applied (`transform`) to the validation folds within an encapsulated `imblearn.pipeline.Pipeline` object.", bold_prefix="3. Encapsulated Pipeline Isolation: ")

    add_heading_2(doc, "5.2 Comprehensive Performance Evaluation Metrics")
    add_body_p(doc, "In imbalanced churn prediction, standard Accuracy is deceptive (e.g., an 85% accurate model that predicts 0 churners is useless). A multi-dimensional metric framework is deployed:")

    metric_headers = ["Evaluation Metric", "Formula / Definition", "Primary Justification in Churn Context"]
    metric_data = [
        ["Precision-Recall AUC (PR-AUC)", "Area under Precision vs. Recall curve", "Focuses strictly on minority churn class performance without being skewed by large true negative counts."],
        ["Recall (Sensitivity)", "TP / (TP + FN)", "Measures fraction of actual churners successfully identified. Crucial because undetected churners (FN) represent lost revenue."],
        ["Precision (Positive Predictive Value)", "TP / (TP + FP)", "Measures accuracy of churn alerts. Prevents wasting retention resources and discounting loyal customers."],
        ["F_Beta Score (β = 2.0)", "(1 + β²)·(P·R) / (β²·P + R)", "Weights Recall twice as heavily as Precision, reflecting the higher financial cost of lost customers relative to outreach cost."],
        ["ROC-AUC", "Area under TPR vs. FPR curve", "Measures overall ranking capability across all possible classification thresholds."],
        ["Brier Score (Calibration Loss)", "(1/N) ∑ (p_i - y_i)²", "Measures reliability of predicted probabilities. Critical for automated risk tiering."],
        ["Cumulative Gains & Lift @ 20%", "Lift = (Churn in Top 20% / Total Churn) / 0.20", "Quantifies how much better the model performs compared to random customer outreach."]
    ]
    tbl_metrics = doc.add_table(rows=1, cols=3)
    style_table(doc, tbl_metrics, [1.8, 1.9, 2.8], metric_headers, metric_data)

    add_heading_2(doc, "5.3 Financial Cost-Benefit Decision Matrix")
    add_body_p(doc, "Rather than relying on an arbitrary 0.50 threshold, the decision threshold θ* is optimized by maximizing expected net financial utility U(θ):")
    add_bullet(doc, "Saved customer lifetime value minus retention campaign cost: V_retained - C_campaign ≈ +$850.", bold_prefix="True Positive (TP) Payoff: ")
    add_bullet(doc, "Cost of outreach/discount provided to a customer who would not have churned: -C_campaign ≈ -$50.", bold_prefix="False Positive (FP) Cost: ")
    add_bullet(doc, "Unmitigated loss of customer lifetime value: -LTV ≈ -$1,200.", bold_prefix="False Negative (FN) Cost: ")
    add_bullet(doc, "Retained customer with zero intervention cost: $0.", bold_prefix="True Negative (TN) Payoff: ")
    add_body_p(doc, "Expected Financial Utility Formulation:")
    add_callout(doc, "Expected Net Profit Optimization Equation",
                "Max Net Utility(θ) = N · [ P(TP|θ) · ($850) + P(FP|θ) · (-$50) + P(FN|θ) · (-$1,200) + P(TN|θ) · ($0) ]\n\nOptimizing this curve across threshold values θ ∈ [0.01, 0.99] determines the maximum ROI operating point (typically θ* ≈ 0.28 to 0.35 in subscription churn).")

    add_heading_2(doc, "5.4 Model Explainability & Interpretability (SHAP / LIME)")
    add_body_p(doc, "To empower customer success teams and meet regulatory transparency standards, TreeSHAP (SHapley Additive exPlanations) is integrated directly into the inference layer:")
    add_bullet(doc, "SHAP feature importance bar charts and beeswarm summary plots identify the macro drivers of enterprise churn (e.g., ticket resolution latency, contract expiration proximity).", bold_prefix="Global Interpretability: ")
    add_bullet(doc, "For every individual high-risk alert, the API generates a localized waterfall plot showing exactly which factors pushed that specific account's risk score above threshold θ*.", bold_prefix="Local Customer Explainability: ")

    # -------------------------------------------------------------
    # SECTION 6: MLOPS PRODUCTION DEPLOYMENT & MAINTENANCE
    # -------------------------------------------------------------
    add_heading_1(doc, "6. MLOps Production Deployment & Maintenance Lifecycle")
    add_body_p(doc, "A machine learning model creates business value only when reliably deployed, monitored, and maintained in production. Figure 4 details the serving architecture and automated drift feedback loops.")

    add_figure(doc, "mlops_deployment_lifecycle.png", "Production Inference, Continuous Observability, and Automated Retraining Architecture.")

    add_heading_2(doc, "6.1 Serving Architecture & Containerization")
    add_bullet(doc, "A high-performance asynchronous REST API built with FastAPI, validating input payloads via Pydantic schemas. Supports batch and single-record predictions with P99 latency < 25ms.", bold_prefix="Real-Time Inference Service: ")
    add_bullet(doc, "Apache Airflow scheduled pipeline running nightly batch scoring on the entire active subscriber base, streaming risk scores directly into Salesforce CRM and Snowflake data warehouse.", bold_prefix="Batch Scoring Worker: ")
    add_bullet(doc, "Models are exported to ONNX (Open Neural Network Exchange) format and serialized via MLflow Model Registry. Services are containerized with Docker and orchestrated via Kubernetes (EKS/GKE).", bold_prefix="Containerization & Registry: ")

    add_heading_2(doc, "6.2 Drift Detection & Continuous Observability")
    add_body_p(doc, "Over time, changes in customer behavior, pricing models, or external market conditions cause model degradation. The observability stack continuously monitors three drift vectors:")
    
    drift_headers = ["Drift Category", "Detection Methodology", "Alert Threshold", "Mitigation Action"]
    drift_data = [
        ["Data Drift (Covariate Shift)", "Population Stability Index (PSI) & 2-Sample Kolmogorov-Smirnov Test on feature distributions.", "PSI > 0.20 on top features; KS p-value < 0.01", "Trigger automated feature pipeline audit; check upstream data ingestion integrity."],
        ["Concept Drift (P(y|X) Shift)", "Tracking rolling PR-AUC and Brier score against realized ground-truth churn events (60-day lag).", "PR-AUC drop > 8% from baseline", "Trigger automated retraining pipeline on latest 12-month rolling data window."],
        ["Operational Drift", "Prometheus & Grafana tracking API request throughput, P95 latency, and memory utilization.", "P95 Latency > 100ms; Error Rate > 0.5%", "Kubernetes horizontal pod autoscaling (HPA) and worker node provisioning."]
    ]
    tbl_drift = doc.add_table(rows=1, cols=4)
    style_table(doc, tbl_drift, [1.5, 2.0, 1.4, 1.6], drift_headers, drift_data)

    add_heading_2(doc, "6.3 Continuous Training (CI/CD/CT) & Deployment Strategy")
    add_bullet(doc, "When drift thresholds are exceeded or on a scheduled monthly cadence, an Airflow DAG spins up a cloud training container, executes nested CV, verifies performance against the champion model, and registers new candidate weights.", bold_prefix="Automated Continuous Retraining (CT): ")
    add_bullet(doc, "New models are deployed via Canary routing (90% traffic to Champion, 10% to Challenger). Once latency, calibration, and stability metrics pass automated smoke tests, traffic shifts to 100%.", bold_prefix="Canary / Blue-Green Deployment: ")
    add_bullet(doc, "Instant zero-downtime rollback to previous artifact version if error rate exceeds 0.1% or unexpected score distribution anomalies occur.", bold_prefix="Automated Rollback Safeguard: ")

    # -------------------------------------------------------------
    # SECTION 7: RISK MANAGEMENT & 35-HOUR IMPLEMENTATION ROADMAP
    # -------------------------------------------------------------
    add_heading_1(doc, "7. Risk Management, Governance & Implementation Roadmap")
    
    add_heading_2(doc, "7.1 Ethical Considerations, Fairness & Data Privacy")
    add_bullet(doc, "Protected demographic attributes (age, gender, ethnicity) are strictly excluded from the feature space. Disparate impact ratios and equal opportunity difference metrics are audited across customer segments to ensure equitable service treatment.", bold_prefix="Algorithmic Fairness: ")
    add_bullet(doc, "All telemetry pipelines comply with GDPR and CCPA standards. Customer identifiers are pseudonymized using SHA-256 salted hashing before feature store ingestion.", bold_prefix="Privacy & Anonymization: ")

    add_heading_2(doc, "7.2 Detailed 35-Hour Implementation Work Breakdown Structure")
    add_body_p(doc, "The following phased work breakdown structure details the 30–35 hours required to execute this comprehensive machine learning implementation:")

    roadmap_headers = ["Phase & Focus Area", "Key Engineering Activities", "Allocated Hours", "Key Deliverables"]
    roadmap_data = [
        ["Phase 1: Inception & ETL", "Define business KPIs; build SQL ingestion extractors; Great Expectations schema validation.", "4 – 5 Hours", "Cleaned data lake tables, schema specification, initial profiling report."],
        ["Phase 2: Preprocessing & Features", "Design scikit-learn Pipeline; imputations; Yeo-Johnson transforms; behavioral feature creation; SMOTE-NC.", "6 – 7 Hours", "Reusable preprocessing pipeline artifact, engineered feature store."],
        ["Phase 3: Model Benchmarking", "Implement baseline Logistic Regression, Random Forest, LightGBM, CatBoost, TabNet in Python.", "6 – 7 Hours", "Candidate model benchmark results table, initial loss curves."],
        ["Phase 4: HPO & Nested CV", "Setup Optuna Bayesian study (150 trials); Nested 5-fold CV; Cost-benefit utility curve thresholding.", "6 – 7 Hours", "Optimized champion weights, hyperparameter convergence plots, utility curve."],
        ["Phase 5: Explainability & Governance", "Generate TreeSHAP summary, waterfall, and dependence plots; fairness audit; documentation.", "3 – 4 Hours", "Explainability dashboard module, algorithmic fairness audit report."],
        ["Phase 6: Deployment & MLOps", "Build FastAPI microservice; Dockerfile containerization; setup Evidently drift monitors and Airflow DAGs.", "5 – 6 Hours", "Working Docker container, REST API endpoints, CI/CD pipeline, MLOps runbook."],
        ["Total Workload", "Full End-to-End Enterprise Implementation", "30 – 35 Hours", "Production-Ready ML Model Package & Documentation"]
    ]
    tbl_road = doc.add_table(rows=1, cols=4)
    style_table(doc, tbl_road, [1.4, 2.3, 1.1, 1.7], roadmap_headers, roadmap_data)

    # -------------------------------------------------------------
    # SECTION 8: CONCLUSION & PROJECT CHECKLIST
    # -------------------------------------------------------------
    add_heading_1(doc, "8. Conclusion & Implementation Checklist")
    add_body_p(doc, "This plan establishes a mathematically sound, reproducible, and commercially impactful blueprint for developing and deploying a machine learning churn prediction engine. By adhering to strict train-test isolation, Bayesian optimization, cost-weighted threshold calibration, and continuous MLOps monitoring, the organization ensures robust model generalization and significant revenue preservation.")

    add_callout(doc, "Phase-Wise Execution Checklist",
                "✓ Problem formulation with concrete financial utility functions.\n"
                "✓ Zero-leakage scikit-learn preprocessing and feature engineering pipeline.\n"
                "✓ Candidate benchmarking across linear, ensemble, and boosted architectures.\n"
                "✓ Optuna Bayesian hyperparameter optimization with median pruning.\n"
                "✓ Stratified Nested K-Fold Cross-Validation with temporal holdout verification.\n"
                "✓ Multidimensional evaluation (PR-AUC, Recall@20, F2-Score, Brier Calibration).\n"
                "✓ Local and Global TreeSHAP interpretability and fairness audits.\n"
                "✓ Containerized FastAPI serving architecture with sub-25ms response time.\n"
                "✓ Continuous drift detection (PSI / KS) and automated CI/CD/CT retraining loops.")

    # Save document
    doc.save(doc_path)
    print(f"Word Document successfully created at: {doc_path}")

if __name__ == "__main__":
    build_document()
