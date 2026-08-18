"""
DOCX Report Generator for Week 2: Exploratory Data Analysis and Visualization
Generates a comprehensive, publication-grade Word Document with embedded figures,
tables, styled callout boxes, code snippets, and in-depth analytical commentary.
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FIGURES_DIR = os.path.join(BASE_DIR, 'figures')
OUTPUT_DOCX = os.path.join(BASE_DIR, 'Week_2_EDA_and_Visualization_Report.docx')


def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner margins (padding) for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def add_callout_box(doc, text, title="KEY TAKEAWAY", border_color="1F4E78", bg_color="F2F7FA"):
    """Create a styled callout / alert box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Border: thick left border, no top/bottom/right border
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r_title = p.add_run(f"[{title}] ")
    r_title.bold = True
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(10)
    r_text.font.italic = True
    r_text.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_code_block(doc, code_str, language="Python"):
    """Create a styled syntax code block."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F8F9FA")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'<w:left w:val="single" w:sz="18" w:space="0" w:color="4A90E2"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    
    r = p.add_run(code_str.strip())
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x2B, 0x3A, 0x42)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def format_table_headers(table, col_widths=None):
    """Apply consistent corporate header styling to tables."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_row = table.rows[0]
    
    # Repeat header row on new pages
    trPr = hdr_row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    
    for idx, cell in enumerate(hdr_row.cells):
        set_cell_background(cell, "1F4E78")
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
    # Style data rows
    for r_idx, row in enumerate(table.rows[1:], start=1):
        bg = "F9FBFD" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, cell in enumerate(row.cells):
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    
    # Apply column widths if provided
    if col_widths:
        for row in table.rows:
            for c_idx, w in enumerate(col_widths):
                if c_idx < len(row.cells):
                    row.cells[c_idx].width = Inches(w)


def add_figure_with_caption(doc, img_filename, fig_num, title, width_inches=6.2):
    """Embed an image with a centered caption."""
    img_path = os.path.join(FIGURES_DIR, img_filename)
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(4)
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(width_inches))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(2)
        p_cap.paragraph_format.space_after = Pt(10)
        
        r_fnum = p_cap.add_run(f"Figure {fig_num}: ")
        r_fnum.bold = True
        r_fnum.font.name = 'Calibri'
        r_fnum.font.size = Pt(9.5)
        r_fnum.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        
        r_title = p_cap.add_run(title)
        r_title.italic = True
        r_title.font.name = 'Calibri'
        r_title.font.size = Pt(9.5)
        r_title.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    else:
        p_err = doc.add_paragraph(f"[Image File Not Found: {img_filename}]")
        p_err.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


def build_report():
    print("Building Week 2 EDA & Visualization Word Report...")
    doc = docx.Document()
    
    # Page setup - 0.8 inch margins
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        
    # Document Title / Cover Header
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    r_prog = title_p.add_run("VIRTUAL DATA SCIENCE WITH PYTHON TRAINEE PROGRAM\n")
    r_prog.font.name = 'Calibri'
    r_prog.font.size = Pt(11)
    r_prog.font.bold = True
    r_prog.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    
    r_main_title = title_p.add_run("Week 2: Exploratory Data Analysis & Visualization Report\n")
    r_main_title.font.name = 'Calibri'
    r_main_title.font.size = Pt(22)
    r_main_title.font.bold = True
    r_main_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    
    r_sub = title_p.add_run("In-Depth Statistical Analysis, Interaction Modeling, and Multi-Factor Visualizations of Healthcare Insurance Expenditures")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(0x4A, 0x60, 0x7A)
    
    # Metadata Block Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Trainee Name / Candidate ID:", "Data Science Trainee - Python Track"),
        ("Program / Track:", "Virtual Data Science with Python Trainee - Week 2 Milestone"),
        ("Primary Dataset:", "Public Healthcare Demographic & Insurance Charges Dataset (1,338 Records)"),
        ("Tools & Libraries Used:", "Python 3.10+, Pandas, NumPy, Matplotlib, Seaborn, SciPy, Scikit-Learn")
    ]
    for idx, (label, val) in enumerate(meta_data):
        cell_lbl, cell_val = meta_table.rows[idx].cells
        set_cell_background(cell_lbl, "F0F4F8")
        set_cell_background(cell_val, "FFFFFF")
        set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
        set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)
        
        p_l = cell_lbl.paragraphs[0]
        r_l = p_l.add_run(label)
        r_l.font.name = 'Calibri'
        r_l.font.bold = True
        r_l.font.size = Pt(9.5)
        r_l.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        
        p_v = cell_val.paragraphs[0]
        r_v = p_v.add_run(val)
        r_v.font.name = 'Calibri'
        r_v.font.size = Pt(9.5)
        r_v.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
    format_table_headers(meta_table, col_widths=[2.4, 4.4])
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # Divider
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(8)
    r_div = p_div.add_run("―" * 65)
    r_div.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    
    # -------------------------------------------------------------------------
    # EXECUTIVE SUMMARY
    # -------------------------------------------------------------------------
    h1 = doc.add_heading("Executive Summary", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    
    doc.add_paragraph(
        "Exploratory Data Analysis (EDA) is the foundational cornerstone of data science, enabling practitioners "
        "to discover underlying data structures, detect anomalies and outliers, test hypotheses, and understand "
        "the relationships between potential predictors and business target metrics. This report details an exhaustive, "
        "end-to-end exploratory analysis conducted on a comprehensive public healthcare and insurance dataset comprising "
        "1,338 policyholders across four distinct geographic regions."
    )
    doc.add_paragraph(
        "Through rigorous univariate, bivariate, and multivariate investigations using Python's scientific ecosystem "
        "(Pandas, Matplotlib, Seaborn, and SciPy), this study evaluates the actuarial and demographic drivers of individual "
        "medical expenditures (`charges`). Key findings demonstrate that while continuous biological aging imparts a steady, "
        "linear escalation on baseline medical costs (~$257 per additional year of age), behavioral smoking status is the single "
        "most dominant financial determinant, increasing average annual costs by over 280% ($32,050.23 for smokers vs. $8,434.27 "
        "for non-smokers)."
    )
    doc.add_paragraph(
        "Crucially, our multivariate interaction analysis reveals a severe non-linear compounding synergy between obesity "
        "(BMI ≥ 30) and smoking: policyholders exhibiting both risk factors face catastrophic annual healthcare costs averaging "
        "$41,557.99—representing 33.94% of total aggregate claims while constituting only 10.8% of the policyholder population. "
        "These empirical insights provide clear pathways for risk stratification, wellness incentive design, and predictive modeling."
    )
    
    add_callout_box(
        doc,
        "Executive Takeaway: Smoking combined with clinical obesity triggers a dramatic non-linear surge in healthcare costs, "
        "quadrupling claim sizes. Logarithmic transformations (log_charges) effectively resolve positive skewness (+1.52 to -0.09) "
        "for downstream regression and predictive modeling.",
        title="CORE EXECUTIVE INSIGHT"
    )
    
    # -------------------------------------------------------------------------
    # SECTION 1: DATASET OVERVIEW & DATA AUDIT
    # -------------------------------------------------------------------------
    h1 = doc.add_heading("1. Dataset Overview & Data Audit", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    
    doc.add_paragraph(
        "The analysis is conducted on the widely recognized Medical Cost Personal Dataset. The dataset captures demographic, "
        "lifestyle, geographic, and family structure attributes for 1,338 individual beneficiaries. Prior to deep analytical "
        "modeling, an exhaustive data audit was executed to ensure data integrity, verify schema consistency, and validate absence "
        "of structural flaws."
    )
    
    # Data Dictionary Table
    doc.add_heading("Data Dictionary & Attribute Specifications", level=2)
    tbl_dict = doc.add_table(rows=8, cols=4)
    tbl_dict.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    dict_headers = ["Attribute", "Data Type", "Domain / Range", "Description & Analytical Role"]
    for i, h in enumerate(dict_headers):
        tbl_dict.rows[0].cells[i].paragraphs[0].add_run(h)
        
    dict_rows = [
        ("age", "Integer (int64)", "18 - 64 years", "Age of primary beneficiary. Proxy for biological aging."),
        ("sex", "Categorical (string)", "female, male", "Insurance contractor gender. Used for demographic parity checks."),
        ("bmi", "Continuous (float64)", "15.96 - 53.13 kg/m²", "Body Mass Index (weight / height²). Indicator of physiological health."),
        ("children", "Integer (int64)", "0 - 5 dependents", "Number of dependent children covered under policy contract."),
        ("smoker", "Categorical (string)", "yes, no", "Smoking status. Major behavioral risk factor variable."),
        ("region", "Categorical (string)", "northeast, northwest, southeast, southwest", "Beneficiary's residential area in the United States."),
        ("charges", "Continuous (float64)", "$1,121.87 - $63,770.43", "Individual medical costs billed to health insurance (Target Variable).")
    ]
    
    for r_i, r_data in enumerate(dict_rows, start=1):
        for c_i, val in enumerate(r_data):
            tbl_dict.rows[r_i].cells[c_i].paragraphs[0].add_run(val)
            
    format_table_headers(tbl_dict, col_widths=[1.2, 1.3, 1.6, 2.7])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # Data Hygiene & Audit Summary
    doc.add_heading("Data Hygiene & Integrity Verification", level=2)
    doc.add_paragraph(
        "A programmatic audit confirmed high baseline data cleanliness:\n"
        "• Missing Value Assessment: 0 missing or null entries detected across all 7 original columns (100% complete records).\n"
        "• Duplicate Record Identification: Exactly 1 duplicate row was detected (a 19-year-old male non-smoker with identical BMI and charges). "
        "This duplicate was documented and preserved as a valid observation of identical baseline characteristics in a large population.\n"
        "• Value Constraints: All continuous attributes fell strictly within clinically plausible biological boundaries (e.g., BMI from 15.96 to 53.13)."
    )
    
    # Summary Statistics Table
    doc.add_heading("Baseline Descriptive Statistics", level=2)
    tbl_stats = doc.add_table(rows=6, cols=8)
    stat_headers = ["Feature", "Count", "Mean", "Std Dev", "Median", "IQR", "Skewness", "Kurtosis"]
    for i, h in enumerate(stat_headers):
        tbl_stats.rows[0].cells[i].paragraphs[0].add_run(h)
        
    stat_rows = [
        ("age", "1,338", "39.21", "14.05", "39.00", "24.00", "0.06", "-1.25"),
        ("bmi", "1,338", "30.66", "6.10", "30.40", "8.40", "0.28", "-0.05"),
        ("children", "1,338", "1.09", "1.21", "1.00", "2.00", "0.94", "0.20"),
        ("charges ($)", "1,338", "13,270.42", "12,110.01", "9,382.03", "11,899.50", "+1.52", "+1.61"),
        ("log_charges", "1,338", "9.10", "0.92", "9.15", "1.34", "-0.09", "-0.64")
    ]
    for r_i, r_data in enumerate(stat_rows, start=1):
        for c_i, val in enumerate(r_data):
            tbl_stats.rows[r_i].cells[c_i].paragraphs[0].add_run(val)
            
    format_table_headers(tbl_stats, col_widths=[1.2, 0.7, 0.9, 0.9, 0.9, 0.8, 0.8, 0.8])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # -------------------------------------------------------------------------
    # SECTION 2: FEATURE ENGINEERING & TRANSFORMATIONS
    # -------------------------------------------------------------------------
    h1 = doc.add_heading("2. Feature Engineering & Data Transformations", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    
    doc.add_paragraph(
        "To enrich the depth of exploratory analysis and uncover hidden multi-way interactions, we engineered four "
        "domain-driven synthetic features:"
    )
    
    doc.add_paragraph(
        "1. WHO BMI Classification (`bmi_category`): Converted raw continuous BMI into standard clinical tiers: "
        "Underweight (<18.5), Normal Weight (18.5 - 24.9), Overweight (25.0 - 29.9), Obese Class I (30.0 - 34.9), "
        "and Obese Class II+ (≥ 35.0).\n"
        "2. Age Cohort Segmentation (`age_group`): Partitioned chronological age into distinct life-cycle brackets: "
        "Young Adult (18-29), Early Career (30-44), Mid Career (45-59), and Senior (60+).\n"
        "3. Target Logarithmic Normalization (`log_charges`): Applied natural log transformation to alleviate extreme positive skewness "
        "in medical charges, reducing skewness from +1.52 to -0.09, establishing near-normality for parametric modeling.\n"
        "4. Multi-Factor Composite Risk Segmentation (`high_risk_segment`): Created a 4-tier combinatorial risk index: "
        "'Healthy Non-Smoker', 'Obese Non-Smoker', 'Smoker Only', and 'Smoker & Obese'."
    )
    
    add_code_block(
        doc,
        "# Feature Engineering and Categorical Transformations in Python\n"
        "df['bmi_category'] = pd.cut(df['bmi'], \n"
        "    bins=[0, 18.5, 24.9, 29.9, 34.9, 100],\n"
        "    labels=['Underweight', 'Normal', 'Overweight', 'Obese Class I', 'Obese Class II+'])\n\n"
        "df['age_group'] = pd.cut(df['age'], \n"
        "    bins=[17, 29, 44, 59, 100],\n"
        "    labels=['Young Adult (18-29)', 'Early Career (30-44)', 'Mid Career (45-59)', 'Senior (60+)'])\n\n"
        "df['log_charges'] = np.log(df['charges'])\n\n"
        "df['high_risk_segment'] = df.apply(\n"
        "    lambda row: 'Smoker & Obese' if (row['smoker'] == 'yes' and row['bmi'] >= 30.0)\n"
        "    else ('Smoker Only' if row['smoker'] == 'yes'\n"
        "          else ('Obese Non-Smoker' if row['bmi'] >= 30.0\n"
        "                else 'Healthy Non-Smoker')),\n"
        "    axis=1\n"
        ")"
    )
    
    # -------------------------------------------------------------------------
    # SECTION 3: UNIVARIATE EXPLORATORY DATA ANALYSIS
    # -------------------------------------------------------------------------
    h1 = doc.add_heading("3. Univariate Exploratory Data Analysis", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    
    doc.add_paragraph(
        "Univariate analysis examines the distribution, central tendency, dispersion, and anomaly patterns of each variable "
        "in isolation. We utilized histogram distributions with Kernel Density Estimation (KDE), reference statistical lines, "
        "and categorical count proportion plots."
    )
    
    # Visual 1
    doc.add_heading("3.1 Target Variable Distribution: Raw vs. Log-Transformed Charges", level=2)
    add_figure_with_caption(
        doc,
        "fig1_univariate_charges_distribution.png",
        fig_num="1",
        title="Univariate Distribution of Medical Charges in Raw USD (left) and Log-Transformed Space (right)."
    )
    doc.add_paragraph(
        "Analysis & Insights:\n"
        "• Raw Charges Distribution: Displays pronounced right-skewness (Skewness = +1.52, Kurtosis = +1.61). The mean ($13,270.42) "
        "is significantly higher than the median ($9,382.03), pulled upward by high-cost catastrophic claims exceeding $40,000.\n"
        "• Multimodality: The KDE curve highlights a primary mode concentrated between $2,000 and $10,000 (standard healthy baseline care), "
        "a secondary peak near $20,000, and a distinct tertiary cluster between $35,000 and $45,000.\n"
        "• Log Transformation Impact: Applying natural logarithm creates an approximately symmetric, bimodal distribution (Skewness = -0.09), "
        "aligning closely with parametric Gaussian assumptions required for linear regression and variance stabilization."
    )
    
    # Visual 2
    doc.add_heading("3.2 Continuous Demographics: Age and Body Mass Index (BMI)", level=2)
    add_figure_with_caption(
        doc,
        "fig2_univariate_age_bmi_distribution.png",
        fig_num="2",
        title="Demographic Histograms and KDE Density for Age (years) and Body Mass Index (BMI in kg/m²)."
    )
    doc.add_paragraph(
        "Analysis & Insights:\n"
        "• Age Distribution: Displays a remarkably uniform spread across the adult lifespan from age 18 to 64 (Mean: 39.2 years, Median: 39.0 years), "
        "with an elevated concentration at age 18-19 (~137 beneficiaries). The low skewness (0.06) confirms balanced age representation.\n"
        "• BMI Distribution: Follows a textbook Gaussian bell curve (Skewness = 0.28, Kurtosis = -0.05) centered at a mean of 30.66 kg/m². "
        "Critically, over 52.8% of the entire sample population exceeds the clinical obesity threshold of BMI = 30.0 kg/m²."
    )
    
    # Visual 3
    doc.add_heading("3.3 Categorical Attribute Distributions", level=2)
    add_figure_with_caption(
        doc,
        "fig3_univariate_categorical_distributions.png",
        fig_num="3",
        title="Frequency Breakdown of Smoking Prevalence, WHO BMI Classes, Geographic Regions, and Dependents."
    )
    doc.add_paragraph(
        "Analysis & Insights:\n"
        "• Smoking Prevalence: 20.5% (274 policyholders) are active smokers, while 79.5% (1,064 policyholders) are non-smokers.\n"
        "• BMI Categorization: Obese Class I and Class II+ represent 29.7% (397) and 23.8% (319) of the population respectively, "
        "meaning only 16.5% (221) fall within the clinically 'Normal' weight category.\n"
        "• Geographic Balance: Policyholders are evenly distributed across the 4 US regions: Southeast (364, 27.2%), Northwest (325, 24.3%), "
        "Southwest (325, 24.3%), and Northeast (324, 24.2%).\n"
        "• Dependent Children: 42.9% (574) have 0 children, 24.2% (324) have 1 child, 17.9% (240) have 2 children, with only 1.3% having 5 children."
    )
    
    # -------------------------------------------------------------------------
    # SECTION 4: BIVARIATE EXPLORATORY DATA ANALYSIS
    # -------------------------------------------------------------------------
    h1 = doc.add_heading("4. Bivariate Exploratory Data Analysis", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    
    doc.add_paragraph(
        "Bivariate analysis investigates how pairs of attributes interact, testing for statistical associations, linear and non-linear "
        "dependencies, and distributional shifts between cohorts."
    )
    
    # Visual 4
    doc.add_heading("4.1 Impact of Smoking Status on Healthcare Expenditures", level=2)
    add_figure_with_caption(
        doc,
        "fig4_bivariate_smoking_impact.png",
        fig_num="4",
        title="Comparative Boxplot with Strip Jitter (left) and Density Violin Plot (right) of Charges by Smoker Status."
    )
    doc.add_paragraph(
        "Analysis & Insights:\n"
        "• Massive Variance Discrepancy: Non-smokers exhibit an average cost of $8,434.27 (Median: $7,345.41, IQR: $7,376.45). "
        "In stark contrast, smokers incur an average cost of $32,050.23 (Median: $34,456.35, IQR: $20,192.96).\n"
        "• Complete Distributional Separation: The interquartile range (IQR) of smokers ($20,826 to $41,019) does not overlap with the "
        "IQR of non-smokers ($4,395 to $11,772), proving that smoking status is an absolute cost multiplier.\n"
        "• Bimodal Violin Density in Smokers: The violin plot for smokers clearly shows two distinct clusters ($15k-$25k and $35k-$50k), "
        "indicating a secondary moderating variable interacting with smoking."
    )
    
    # Visual 5
    doc.add_heading("4.2 Age vs. Charges Moderated by Smoking Status", level=2)
    add_figure_with_caption(
        doc,
        "fig5_bivariate_age_charges_by_smoker.png",
        fig_num="5",
        title="Scatter Plot of Age vs. Medical Charges with Ordinary Least Squares (OLS) Fitted Regression Lines by Smoker Status."
    )
    doc.add_paragraph(
        "Analysis & Insights:\n"
        "• Parallel Baseline Slopes: For non-smokers, charges scale linearly with age at an empirical rate of approximately +$267 per year. "
        "A healthy 20-year-old averages ~$2,500, whereas a healthy 60-year-old averages ~$13,000.\n"
        "• Shifted Smoker Planes: For smokers, the regression trend is elevated by an intercept shift of over $23,000. Furthermore, smokers "
        "split into two parallel tracks: a lower track (smokers with normal/moderate BMI) and an upper track (smokers with high BMI)."
    )
    
    # Visual 6
    doc.add_heading("4.3 BMI vs. Charges: Non-Linear Threshold Interaction", level=2)
    add_figure_with_caption(
        doc,
        "fig6_bivariate_bmi_charges_interaction.png",
        fig_num="6",
        title="Scatter Plot of BMI vs. Medical Charges Demonstrating the Critical Obesity (BMI = 30) Inflection Point."
    )
    doc.add_paragraph(
        "Analysis & Insights:\n"
        "• Non-Smoker Baseline Flatness: For non-smokers (teal points), increasing BMI creates only a mild, gradual increase in claims. "
        "Over 95% of non-smokers remain below $15,000 regardless of whether their BMI is 20 or 45.\n"
        "• The Critical BMI = 30 Bifurcation: For smokers (red markers), BMI < 30 produces charges clustered between $13,000 and $25,000. "
        "However, the moment BMI crosses the clinical obesity threshold of 30.0 kg/m², charges instantly jump into the catastrophic "
        "range of $35,000 to $60,000. This confirms a classic non-linear threshold interaction effect."
    )
    
    # Visual 7
    doc.add_heading("4.4 Regional Variations in Costs and Health Indicators", level=2)
    add_figure_with_caption(
        doc,
        "fig7_bivariate_regional_analysis.png",
        fig_num="7",
        title="Mean Medical Charges with 95% CI (left) and BMI Boxplots (right) Across US Geographic Regions."
    )
    doc.add_paragraph(
        "Analysis & Insights:\n"
        "• Southeast Regional Peak: The Southeast region exhibits the highest mean charges ($14,735.41), significantly outstripping the "
        "Southwest ($12,346.94) and Northwest ($12,417.58).\n"
        "• Underlying Drivers: The right boxplot explains this disparity: the Southeast has the highest median BMI (33.36 kg/m²) and the highest "
        "regional smoking prevalence (25.0%), driving a greater proportion of its population into the catastrophic high-risk quadrant."
    )
    
    # -------------------------------------------------------------------------
    # SECTION 5: MULTIVARIATE EDA & STATISTICAL AGGREGATIONS
    # -------------------------------------------------------------------------
    h1 = doc.add_heading("5. Multivariate Exploratory Data Analysis & Statistical Modeling", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    
    doc.add_paragraph(
        "Multivariate analysis simultaneously evaluates complex correlations, higher-order feature interactions, and conditional "
        "sub-cohort variations across multiple dimensions."
    )
    
    # Visual 8
    doc.add_heading("5.1 Multivariate Correlation Matrix", level=2)
    add_figure_with_caption(
        doc,
        "fig8_multivariate_correlation_heatmap.png",
        fig_num="8",
        title="Pearson Correlation Heatmap of Demographic, Physiological, and Cost Determinants.",
        width_inches=5.8
    )
    doc.add_paragraph(
        "Analysis & Insights:\n"
        "• Smoker vs. Charges (r = 0.787): Smoking status demonstrates the strongest linear correlation with medical charges, followed by Age (r = 0.299) "
        "and BMI (r = 0.198).\n"
        "• Independence of Predictors: The correlation between Age and BMI is near zero (r = 0.109), and between Age and Smoker is r = -0.025. "
        "This confirms low multicollinearity among independent variables, ensuring stable regression coefficients in downstream modeling."
    )
    
    # Visual 9
    doc.add_heading("5.2 Multi-Panel FacetGrid: Age vs. Charges Across BMI Categories", level=2)
    add_figure_with_caption(
        doc,
        "fig9_multivariate_facetgrid_bmi_age_smoker.png",
        fig_num="9",
        title="FacetGrid Decomposition: Scatter Plots of Age vs. Charges Across 5 WHO BMI Classes Stratified by Smoking Behavior."
    )
    doc.add_paragraph(
        "Analysis & Insights:\n"
        "• Underweight / Normal / Overweight Panels: Smoker charges (red) remain in the $15k-$25k tier across ages 18 to 64.\n"
        "• Obese Class I & II+ Panels: The smoker cluster completely shifts into the $35k-$60k tier, showing how the combination of "
        "high BMI and smoking completely alters the financial risk profile."
    )
    
    # Visual 10
    doc.add_heading("5.3 Composite Risk Segmentation Analysis", level=2)
    add_figure_with_caption(
        doc,
        "fig10_multivariate_risk_segmentation.png",
        fig_num="10",
        title="Boxplot Distribution of Annual Medical Charges Across 4 Composite Risk Cohorts."
    )
    
    # Risk Cohort Table
    tbl_risk = doc.add_table(rows=5, cols=5)
    risk_headers = ["Risk Segment Cohort", "Beneficiary Count", "Mean Charges ($)", "Median Charges ($)", "Share of Total Costs (%)"]
    for i, h in enumerate(risk_headers):
        tbl_risk.rows[0].cells[i].paragraphs[0].add_run(h)
        
    risk_rows = [
        ("Healthy Non-Smoker (BMI < 30)", "502 (37.5%)", "$7,977.03", "$6,761.62", "22.55%"),
        ("Obese Non-Smoker (BMI ≥ 30)", "562 (42.0%)", "$8,842.69", "$8,076.05", "27.99%"),
        ("Smoker Only (BMI < 30)", "129 (9.6%)", "$21,363.22", "$20,167.34", "15.52%"),
        ("Smoker & Obese (BMI ≥ 30)", "145 (10.8%)", "$41,557.99", "$40,904.20", "33.94%")
    ]
    for r_i, r_data in enumerate(risk_rows, start=1):
        for c_i, val in enumerate(r_data):
            tbl_risk.rows[r_i].cells[c_i].paragraphs[0].add_run(val)
            
    format_table_headers(tbl_risk, col_widths=[2.0, 1.2, 1.2, 1.2, 1.2])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # Visual 11
    doc.add_heading("5.4 Pairwise Feature Interactions (Pairplot)", level=2)
    add_figure_with_caption(
        doc,
        "fig11_multivariate_pairplot.png",
        fig_num="11",
        title="Comprehensive Pairplot Matrix of Continuous Variables Colored by Smoker Status.",
        width_inches=5.8
    )
    doc.add_paragraph(
        "Analysis & Insights:\n"
        "• Diagonal KDE plots confirm that while Age and BMI have similar distributions across smokers and non-smokers, the target "
        "variable (charges) separates cleanly into distinct bimodal distributions.\n"
        "• Off-diagonal scatter plots emphasize the clean geometric hyperplane separation between risk groups, confirming high separability "
        "for classification and tree-based regression algorithms."
    )
    
    # -------------------------------------------------------------------------
    # SECTION 6: CRITICAL ACTUARIAL, CLINICAL & BUSINESS INSIGHTS
    # -------------------------------------------------------------------------
    h1 = doc.add_heading("6. Critical Actuarial, Clinical & Business Insights", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    
    doc.add_paragraph(
        "Synthesizing our quantitative findings yields three major strategic conclusions for healthcare administrators, "
        "actuaries, and policy underwriters:"
    )
    
    doc.add_paragraph(
        "1. The Synergistic Multiplier Effect (Smoking × Obesity):\n"
        "Individually, smoking adds ~$13,386 in expected annual costs over baseline, while obesity alone adds only ~$865. "
        "However, when both risk factors co-occur, they generate a super-additive cost explosion averaging $41,557.99 per policyholder. "
        "This 10.8% subgroup consumes over one-third (33.94%) of all aggregate insurance disbursements."
    )
    
    doc.add_paragraph(
        "2. Predictable Biological Aging vs. Volatile Behavioral Lifestyle:\n"
        "Biological aging is a stable, deterministic predictor (+$267/year). Underwriting models can safely parameterize age with standard "
        "linear coefficients. Conversely, lifestyle behaviors (smoking, extreme BMI) introduce severe variance spikes that require non-linear "
        "tree models or interaction terms."
    )
    
    doc.add_paragraph(
        "3. Geographic Disparities Stem from Population Health Demographics:\n"
        "The elevated charges observed in the Southeast region are not an artifact of regional medical billing rates alone; they reflect "
        "statistically higher rates of adult obesity (33.36 kg/m² mean) and higher smoking rates (25.0%). Regional premium adjustments "
        "must account for underlying health risk density."
    )
    
    # -------------------------------------------------------------------------
    # SECTION 7: STRATEGIC RECOMMENDATIONS & NEXT STEPS
    # -------------------------------------------------------------------------
    h1 = doc.add_heading("7. Strategic Recommendations & Predictive Modeling Roadmap", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    
    doc.add_paragraph(
        "Based on the empirical evidence gathered during this EDA milestone, the following strategic actions and predictive modeling "
        "steps are recommended:"
    )
    
    doc.add_paragraph(
        "• Dynamic Actuarial Pricing Tiers: Implement non-linear risk rating for policyholders who both smoke and possess a BMI ≥ 30. "
        "Standard additive premium surcharges underprice this cohort, leading to underwriting loss ratios.\n"
        "• Wellness and Smoking Cessation Subsidies: Because smoking cessation reduces expected claim costs from $41,557 to $8,842 in obese "
        "individuals, fully subsidized smoking cessation and weight management programs will deliver massive ROI.\n"
        "• Modeling Recommendations for Machine Learning Pipeline:\n"
        "  a) Target Transformation: Use `log_charges` as the target for Ordinary Least Squares (OLS), Ridge, and Lasso regression to prevent "
        "heteroscedasticity and high residual errors on catastrophic claims.\n"
        "  b) Interaction Features: Explicitly include polynomial and interaction terms such as `bmi * is_smoker` and `age * is_smoker`.\n"
        "  c) Tree-Based Ensembles: Deploy Gradient Boosted Decision Trees (XGBoost / LightGBM) or Random Forests, which natively capture "
        "non-linear threshold cutoffs (e.g., BMI = 30) without manual binning."
    )
    
    # -------------------------------------------------------------------------
    # SECTION 8: FULL PYTHON CODE APPENDIX
    # -------------------------------------------------------------------------
    h1 = doc.add_heading("8. Full Python Code Implementation Appendix", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    
    doc.add_paragraph(
        "Below is the complete, modular Python script used to perform data ingestion, feature engineering, statistical summaries, "
        "and high-resolution figure exports:"
    )
    
    # Read eda_analysis.py code
    eda_code_path = os.path.join(BASE_DIR, 'eda_analysis.py')
    with open(eda_code_path, 'r', encoding='utf-8') as f:
        full_code = f.read()
        
    add_code_block(doc, full_code)
    
    # Save the document
    doc.save(OUTPUT_DOCX)
    print(f"Report generated and saved successfully to: {OUTPUT_DOCX}")


if __name__ == '__main__':
    build_report()
